import base64
import html as html_lib
import json
import os
import re
import shutil
import time
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from urllib.error import HTTPError, URLError
from urllib.parse import parse_qs, quote, unquote, urlencode, urlparse
from urllib.request import Request, urlopen
from zipfile import ZIP_DEFLATED, ZipFile

from lxml import etree
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
STATIC = ROOT / "static"
GENERATED = ROOT / "generated"
OUTPUTS = ROOT.parents[1] / "outputs"
UPLOADS = ROOT / "uploads"
# 面向大众版：素材库只扫描项目内的 materials/ 目录，不再扫描开发者个人桌面，
# 避免把私人资料（简历、实习文件等）注入他人简历。如需内置通用素材，放到 server.py 同级的 materials/ 即可。
PROJECT_DIR = Path(__file__).resolve().parent
MATERIALS_DIR = PROJECT_DIR / "materials"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}
GENERATED.mkdir(exist_ok=True)
OUTPUTS.mkdir(exist_ok=True)
UPLOADS.mkdir(exist_ok=True)

RUNTIME_DEEPSEEK_API_KEY = ""
LLM_KEY_FILE = Path(__file__).resolve().parent / ".llm_key"
RUNTIME_DEEPSEEK_MODEL = ""
LIBRARY_FILE_CACHE = {"time": 0.0, "root": "", "files": []}
TEXT_CACHE = {}


BASE_PROFILE = {
    "name": "候选人",
    "contact": "",
    "education": [],
    "honors": "",
    "experiences": [],
    "competitions": [],
    "research": [["", "", ""], ["", "", ""]],
}


SECTOR_KEYWORDS = {
    "lawfirm": ["律所", "律师", "争议", "诉讼", "仲裁", "并购", "m&a", "尽调", "合规", "知识产权", "商标", "专利", "法律检索", "memo", "英文"],
    "finance": ["金融", "券商", "基金", "投行", "投资", "估值", "财务", "尽调", "行业研究", "资本市场", "ipo", "债券", "wind", "excel", "风险"],
    "soe": ["国企", "央企", "国央企", "法务", "合规", "内控", "政策", "党建", "公文", "综合文字", "合同", "采购", "监管", "政府"],
}

JD_TERMS = {
    "尽调": ["尽职调查", "工商", "股权", "涉诉", "负面信息", "第三方同意"],
    "合规": ["合规", "合规风险", "合规审查", "监管", "审批", "备案", "内控", "风险边界", "风险控制", "合同审核", "制度建设"],
    "数据合规": ["数据合规", "个人信息", "数据出境", "网络安全", "算法", "生成式人工智能", "aigc", "ai出海", "AI 出海", "技术出口管制"],
    "文书": ["法律文书", "文书起草", "报告", "memo", "邮件", "公文", "材料撰写"],
    "英文": ["英文", "双语", "跨境", "境外客户", "英语"],
    "争议": ["争议", "诉讼", "仲裁", "证据", "裁判规则"],
    "金融": ["金融", "投融资", "债务重组", "投资退出", "资本市场", "金融监管", "结构性产品", "final terms", "equity-linked", "autocallable", "hkex", "外资银行", "资管", "衍生品"],
    "基金": ["基金", "私募", "股权投资", "投决", "投资建议书", "绿色投资", "风控", "风险评估", "募投管退", "投后", "估值", "行业研究", "产业研究"],
    "知识产权": ["知识产权", "商标", "专利", "不正当竞争", "装潢"],
    "研究": ["研究", "检索", "政策", "行业", "分析"],
}

CONTEXT_SIGNALS = {
    "数据合规": ["数据合规", "个人信息", "个人信息保护", "数据出境", "数据安全", "网络安全", "算法", "生成式人工智能", "aigc", "ai出海", "技术出口管制", "隐私政策", "pipl"],
    "金融监管": ["金融监管", "final terms", "结构性产品", "结构性票据", "equity-linked", "autocall", "autocallable", "hkex", "外资银行", "衍生品", "票据发行", "calculation agent", "selling restriction", "资管新规", "银行监管"],
    "金融投资": ["投资", "投融资", "基金", "私募", "股权投资", "投资建议书", "投决", "投委会", "募投管退", "投后", "投资退出", "行业研究", "产业研究", "估值", "风控", "风险评估", "尽调", "pe", "vc", "资产管理", "项目收益", "现金流", "财务模型"],
    "绿色投资": ["绿色投资", "新能源", "风能", "地热", "saf", "管网", "绿色产业", "碳", "环保", "风电", "光伏", "储能"],
    "普通合规": ["合规", "合规审查", "合规风险", "监管", "审批", "备案", "内控", "风险控制", "风险识别", "制度", "流程", "合同审核", "合同审查", "法律风险", "经营合规", "反垄断", "反商业贿赂"],
    "国央企法务": ["国企", "央企", "国央企", "政府", "事业单位", "公文", "综合文字", "政策研究", "党建", "采购", "招投标", "内部审批", "制度建设"],
    "知识产权": ["知识产权", "商标", "专利", "著作权", "不正当竞争", "商业秘密", "装潢", "侵权", "无效", "异议", "撤三"],
    "争议解决": ["争议", "诉讼", "仲裁", "证据", "庭审", "裁判规则", "案件", "执行", "抗诉", "纠纷", "法律检索"],
    "英文": ["英文", "英语", "双语", "跨境", "境外", "海外", "邮件", "翻译", "english"],
    "文书写作": ["文书", "报告", "memo", "备忘录", "公文", "材料", "起草", "撰写", "legal research"],
}

MATERIAL_KEYWORDS = {
    "AI/数据合规": ["AI Lab", "AI出海", "AI Lab相关", "技术出口管制", "生成式人工智能", "互联网信息服务", "数据", "出口管制"],
    "金融监管/Final Terms": ["final terms", "structured products", "equity-linked notes", "autocallable", "calculation agent", "hkex", "bank of america", "merrill lynch"],
    "SEP/专利": ["SEP", "专利", "侵权赔偿"],
    "商品包装装潢": ["商品包装", "装潢", "反不正当竞争", "显著性"],
    "金融/私募基金": ["绿色私募", "私募", "基金", "投资建议书", "投决", "风险评估", "财务尽职调查", "股权投资", "final terms", "structured products", "equity-linked notes", "hkex", "bank of america"],
    "商事争议": ["酒店管理", "职业放贷", "不当得利", "票据纠纷", "山东物流"],
    "商标": ["商标", "知产"],
    "模拟立法/AIGC": ["模拟立法", "全国法科学生模拟立法大赛", "生成式人工智能", "aigc", "训练数据", "算法治理", "人工智能立法"],
}

SUPPORTED_LIBRARY_SUFFIXES = {".docx", ".pdf", ".txt", ".md", ".rtf"}
QUERY_STOPWORDS = {
    "简历", "修改", "要求", "内容", "根据", "针对", "岗位", "工作", "实习", "经历", "可以", "需要",
    "这个", "那个", "里面", "文件", "资料", "生成", "增加", "丰富", "突出", "强调", "不要", "不用",
    "删除", "删掉", "部分", "方向", "能力", "相关", "进行", "负责", "协助",
}
BROAD_SCAN_TERMS = {
    "研究", "法律", "法务", "合规", "文书", "材料", "撰写", "审核", "合同", "岗位", "简历", "实习",
    "分析", "负责", "协助", "风险", "工作", "内容", "修改", "要求", "检索", "最新", "语境", "针对性",
}
SHORT_IMPORTANT_TERMS = {"ai", "cc", "ma", "m&a", "sep", "ldr", "hkex", "pe", "vc", "jd"}
WEB_WEAK_TERMS = {"ai", "jd", "cc", "监管", "法律实务", "服务", "管理", "生成", "作成", "数据", "个人", "信息", "个人信息", "合规", "法务"}
WEB_SHORT_ALLOWED = {"hkex", "finalterms", "autocallable", "equitylinked", "calculationagent", "sep", "pipl"}


def extract_text_from_path(path, limit=4000):
    try:
        stat = path.stat()
        cache_key = (str(path), stat.st_mtime, stat.st_size, limit)
        if cache_key in TEXT_CACHE:
            return TEXT_CACHE[cache_key]
    except OSError:
        cache_key = None
    suffix = path.suffix.lower()
    try:
        if suffix in {".txt", ".md", ".rtf"}:
            text = path.read_text(errors="ignore")[:limit]
            if cache_key:
                TEXT_CACHE[cache_key] = text
            return text
        if suffix == ".docx":
            doc = Document(path)
            para_parts = [p.text for p in doc.paragraphs if p.text.strip()]
            table_parts = []
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        table_parts.append("\t".join(cells))
            # Resumes usually keep name/contact in the first table at the top;
            # move it to the front so the header parser sees it.
            parts = (table_parts[:1] + para_parts + table_parts[1:]) if table_parts else para_parts
            text = "\n".join(parts)[:limit]
            if cache_key:
                TEXT_CACHE[cache_key] = text
            return text
        if suffix == ".pdf":
            import pdfplumber

            chunks = []
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages[:4]:
                    chunks.append(page.extract_text() or "")
                    if sum(len(c) for c in chunks) > limit:
                        break
            text = "\n".join(chunks)[:limit]
            if cache_key:
                TEXT_CACHE[cache_key] = text
            return text
    except Exception:
        return ""
    return ""


def compact_text(text, limit=520):
    return re.sub(r"\s+", " ", text or "").strip()[:limit]


def query_terms_from_text(*texts):
    raw = "\n".join(text or "" for text in texts)
    normalized = normalize_text(raw)
    terms = []
    for groups in (CONTEXT_SIGNALS, JD_TERMS, MATERIAL_KEYWORDS):
        for words in groups.values():
            for word in words:
                key = normalize_text(word)
                if key and key in normalized:
                    terms.append(word)
    for token in re.findall(r"[A-Za-z][A-Za-z0-9_\-/]{1,}|[0-9]{2,}|[\u4e00-\u9fff]{2,}", raw):
        token = token.strip()
        if token.lower() in QUERY_STOPWORDS or token in QUERY_STOPWORDS:
            continue
        if re.fullmatch(r"[\u4e00-\u9fff]{9,}", token):
            for length in (8, 6, 4):
                for i in range(0, max(0, len(token) - length + 1), max(1, length - 2)):
                    piece = token[i : i + length]
                    if piece not in QUERY_STOPWORDS:
                        terms.append(piece)
        else:
            terms.append(token)
    seen = set()
    unique = []
    for term in terms:
        key = normalize_text(term)
        if not key or key in seen or key in QUERY_STOPWORDS:
            continue
        seen.add(key)
        unique.append(term)
    return unique[:80]


def useful_scan_term(term):
    key = normalize_text(term)
    if not key or key in BROAD_SCAN_TERMS or key in QUERY_STOPWORDS:
        return False
    if key in SHORT_IMPORTANT_TERMS:
        return True
    if re.fullmatch(r"20\d{2}|\d+", key):
        return False
    return len(key) >= 3


def safe_relative(path, root):
    try:
        return str(path.relative_to(root))
    except ValueError:
        return str(path)


def is_scannable_material(path):
    if not path.is_file():
        return False
    if path.name.startswith(".") or path.name.startswith("~") or path.name.startswith("~$"):
        return False
    if path.suffix.lower() not in SUPPORTED_LIBRARY_SUFFIXES:
        return False
    try:
        return path.stat().st_size <= 18 * 1024 * 1024
    except OSError:
        return False


def scan_internship_library(jd="", custom=""):
    root = MATERIALS_DIR
    if not root.exists():
        return {"count": 0, "matches": [], "summary": "未配置共享素材库。", "snippets": []}
    query = normalize_text(f"{jd}\n{custom}")
    query_terms = query_terms_from_text(jd, custom)
    now = time.time()
    if LIBRARY_FILE_CACHE["root"] == str(root) and now - LIBRARY_FILE_CACHE["time"] < 90:
        files = LIBRARY_FILE_CACHE["files"]
    else:
        files = []
        for path in root.rglob("*"):
            if is_scannable_material(path):
                files.append(path)
            if len(files) >= 1600:
                break
        LIBRARY_FILE_CACHE.update({"time": now, "root": str(root), "files": files})
    matches = []
    text_reads = 0
    for path in files:
        rel = safe_relative(path, root)
        name = rel.lower()
        name_normalized = normalize_text(rel)
        in_internship_folder = "实习" in Path(rel).parts
        score = 0
        tags = []
        for tag, words in MATERIAL_KEYWORDS.items():
            file_hit = any(w.lower().replace(" ", "") in name.replace(" ", "") for w in words)
            query_hit = any(w.lower().replace(" ", "") in query for w in words)
            if file_hit:
                score += 3 + (4 if query_hit else 0)
                tags.append(tag)
        for term in query_terms:
            key = normalize_text(term)
            if useful_scan_term(term) and key in name_normalized:
                score += 3
                tags.append("文件名命中")
        if in_internship_folder:
            score += 1
        snippet = ""
        generic_material_name = any(word in name_normalized for word in ["自我介绍", "简历面", "实习资料", "面试", "resume", "cv"])
        suffix = path.suffix.lower()
        text_friendly = suffix in {".docx", ".txt", ".md", ".rtf"}
        explicit_material_hit = any(tag not in {"文件名命中"} for tag in tags)
        should_read_text = (
            (text_friendly and (score > 0 or in_internship_folder or generic_material_name))
            or (suffix == ".pdf" and explicit_material_hit and score >= 6)
        )
        if should_read_text and text_reads < 90:
            text_reads += 1
            snippet = compact_text(extract_text_from_path(path, limit=1800), 680)
            content = normalize_text(snippet)
            content_hits = []
            for term in query_terms:
                key = normalize_text(term)
                if useful_scan_term(term) and key in content:
                    content_hits.append(term)
            if content_hits:
                score += min(18, len(content_hits) * 3)
                tags.append("正文命中")
        if score:
            unique_tags = []
            for tag in tags:
                if tag not in unique_tags:
                    unique_tags.append(tag)
            matches.append({"name": rel, "tags": unique_tags[:4], "score": score, "snippet": snippet})
    matches.sort(key=lambda item: (-item["score"], item["name"]))
    matches = matches[:18]
    summary = "；".join(f"{m['name']}（{ '、'.join(m['tags']) }）" for m in matches[:8])
    snippets = [
        {"name": item["name"], "tags": item["tags"], "snippet": item.get("snippet", "")}
        for item in matches[:8]
        if item.get("snippet")
    ]
    base_summary = summary or "暂未命中特定资料；可在修改要求中写明文件名、业务词或项目方向。"
    return {"count": len(files), "matches": matches, "summary": base_summary, "snippets": snippets, "root": str(root)}


def save_uploaded_file(payload):
    filename = re.sub(r"[^A-Za-z0-9._\-\u4e00-\u9fff]", "_", payload.get("name", f"upload_{int(time.time())}"))
    data = payload.get("data", "")
    if "," in data:
        data = data.split(",", 1)[1]
    raw = base64.b64decode(data)
    path = UPLOADS / f"{int(time.time() * 1000)}_{filename}"
    path.write_bytes(raw)
    text = extract_text_from_path(path)
    if not text and path.suffix.lower() in {".png", ".jpg", ".jpeg", ".webp", ".heic"}:
        text = f"[图片材料：{filename}] 暂未做 OCR，可在修改要求中补充图片要点。"
    return {"name": filename, "size": len(raw), "text": text[:5000], "stored": str(path)}


def normalize_text(text):
    return (text or "").lower().replace(" ", "")


def has_any(text, words):
    normalized = normalize_text(text)
    return any(word.lower().replace(" ", "") in normalized for word in words)


def hit_words(text, words):
    normalized = normalize_text(text)
    hits = []
    for word in words:
        key = word.lower().replace(" ", "")
        if key and key in normalized:
            hits.append(word)
    return hits


def analyze_jd_context(jd="", custom=""):
    signal_text = f"{jd or ''}\n{custom or ''}"
    signals = {}
    for label, words in CONTEXT_SIGNALS.items():
        hits = hit_words(signal_text, words)
        signals[label] = {
            "hits": hits,
            "score": len(hits),
            "jdHits": hit_words(jd or "", words),
            "customHits": hit_words(custom or "", words),
        }
    return signals


def signal_active(signals, label):
    return bool(signals.get(label, {}).get("hits"))


def signal_summary(signals):
    parts = []
    for label, item in signals.items():
        hits = item.get("hits") or []
        if hits:
            parts.append(f"{label}：{ '、'.join(hits[:6]) }")
    return "；".join(parts)


def custom_says_no(custom, labels):
    normalized = normalize_text(custom)
    for label in labels:
        key = label.lower().replace(" ", "")
        for prefix in ["不要", "不要写", "不写", "别写", "别再写", "删除", "删掉", "弱化", "不要再写", "不需要"]:
            if f"{prefix}{key}" in normalized:
                return True
    return False


def custom_says_yes(custom, labels):
    normalized = normalize_text(custom)
    for label in labels:
        key = label.lower().replace(" ", "")
        for prefix in ["增加", "加入", "突出", "强调", "保留", "替换为", "替换成", "改为", "改成", "丰富", "补充", "扩写", "展开", "细化"]:
            if f"{prefix}{key}" in normalized or (prefix in normalized and key in normalized):
                return True
    return False


def custom_mentions(custom, labels):
    normalized = normalize_text(custom)
    return any(normalize_text(label) in normalized for label in labels)


def apply_manual_content_requests(resume, custom):
    if not custom.strip():
        return resume
    opinions = list(resume.get("opinions", []))
    terms = list(resume.get("terms", []))
    research = list(resume.get("research", []))
    competitions = list(resume.get("competitions", []))
    changed = []

    wants_legislation = (
        custom_says_yes(custom, ["模拟立法", "全国法科学生模拟立法大赛", "AIGC", "生成式人工智能"])
        or (
            custom_mentions(custom, ["模拟立法", "全国法科学生模拟立法大赛"])
            and custom_mentions(custom, ["丰富", "补充", "扩写", "展开", "细化", "增加"])
        )
    )
    if wants_legislation:
        enriched = (
            "围绕生成式人工智能服务管理开展规则设计，研究训练数据来源合法性、知识产权保护、"
            "生成内容安全、服务提供者义务及监管备案机制；参与草案条文与立法说明撰写，"
            "将技术风险转化为可操作的义务配置和责任边界。"
        )
        found = False
        next_research = []
        for title, date, detail in research:
            if "模拟立法" in title:
                next_research.append((title, date, enriched))
                found = True
            else:
                next_research.append((title, date, detail))
        if not found:
            next_research.append(("2023全国法科学生模拟立法大赛 | 项目核心成员", "2023.07 – 2023.09", enriched))
        research = next_research[:3]
        if "模拟立法" not in terms:
            terms.append("模拟立法")
        changed.append("已按你的要求扩写 2023 全国法科学生模拟立法大赛，突出 AIGC 规则设计、训练数据合法性、知识产权、内容安全、服务提供者义务和监管备案机制。")

    if custom_says_no(custom, ["竞赛", "模拟法庭"]) and not wants_legislation:
        competitions = competitions[:1]
        changed.append("已按要求压缩竞赛经历，只保留最匹配的一段。")

    if changed:
        opinions.insert(0, {"title": "手写要求已写入正文", "body": " ".join(changed)})
        resume = {**resume, "terms": terms, "research": research, "competitions": competitions, "opinions": opinions}
    return resume


SECTION_ALIASES = {
    "education": ["教育背景", "教育经历", "学习经历", "education"],
    "experience": ["实习经历", "工作经历", "实践经历", "职业经历", "professional experience", "internship experience", "work experience"],
    "project": ["项目经历", "项目经验", "社会实践", "校园经历", "活动经历", "project experience"],
    "competition": ["竞赛经历", "比赛经历", "获奖经历", "competition", "awards"],
    "research": ["科研经历", "论文", "研究经历", "research", "publications"],
    "skills": ["技能", "技能与兴趣", "其他", "证书", "语言", "skills"],
}


def normalize_resume_lines(text):
    lines = []
    for line in re.split(r"[\r\n]+", text or ""):
        line = re.sub(r"\s+", " ", line).strip()
        if line:
            lines.append(line)
    return lines


def detect_section(line):
    normalized = normalize_text(line)
    trimmed = re.sub(r"[\s:：｜|/]+", "", line).lower()
    if len(trimmed) > 28:
        return ""
    for key, names in SECTION_ALIASES.items():
        if any(normalize_text(name) == normalized or normalize_text(name) in normalized for name in names):
            return key
    return ""


def split_resume_sections(lines):
    sections = {"header": []}
    current = "header"
    for line in lines:
        section = detect_section(line)
        if section:
            current = section
            sections.setdefault(current, [])
            continue
        sections.setdefault(current, []).append(line)
    return sections


DATE_PATTERN = re.compile(
    r"(20\d{2}[./年-]?\s?\d{0,2}\s*(?:[–—\-至到~]+\s*(?:20\d{2}[./年-]?\s?\d{0,2}|至今|现在|Present|Current))?|20\d{2}[./年-]\s?\d{1,2})",
    re.I,
)


def extract_date_from_line(line):
    match = DATE_PATTERN.search(line or "")
    if not match:
        return "", line
    date = match.group(1).strip()
    rest = (line[: match.start()] + " " + line[match.end() :]).strip(" |｜,，;；")
    return date, re.sub(r"\s+", " ", rest).strip()


def clean_bullet_text(line):
    text = re.sub(r"^[•·▪●◦\-\*\uF0B7\s]+", "", line or "").strip()
    text = re.sub(r"^(负责|参与|协助|主要工作|工作内容)[:：]\s*", "", text)
    return text.strip("；; ")


def choose_label(text):
    normalized = normalize_text(text)
    if has_any(normalized, ["数据", "个人信息", "算法", "ai", "aigc", "出境", "出口管制"]):
        return "数据与技术合规："
    if has_any(normalized, ["金融", "基金", "投资", "投融资", "尽调", "估值", "财务", "行业研究"]):
        return "金融与交易支持："
    if has_any(normalized, ["合规", "监管", "风险", "内控", "审查", "合同"]):
        return "合规与风险识别："
    if has_any(normalized, ["检索", "研究", "政策", "分析", "调研"]):
        return "研究分析："
    if has_any(normalized, ["起草", "撰写", "报告", "memo", "备忘录", "文书", "邮件"]):
        return "文书写作："
    if has_any(normalized, ["英文", "英语", "翻译", "跨境"]):
        return "英文与跨境支持："
    return "工作内容："


def score_text_for_jd(text, jd, custom, terms):
    haystack = normalize_text(text)
    score = 0
    for term in terms:
        if normalize_text(term) in haystack:
            score += 4
    for group in CONTEXT_SIGNALS.values():
        for word in group:
            key = normalize_text(word)
            if key and key in normalize_text(f"{jd}\n{custom}") and key in haystack:
                score += 3
    return score


def parse_profile_from_external(lines, sections):
    head = sections.get("header", [])[:8]
    all_text = "\n".join(lines[:18])
    email = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", all_text)
    phone = re.search(r"(?:\+?86[-\s]?)?1[3-9]\d{9}|\+?\d[\d\s\-()]{7,}\d", all_text)
    links = re.findall(r"(?:https?://|www\.)[^\s；;，,]+", all_text)
    name = ""
    for line in head:
        candidate = re.sub(r"(电话|手机|邮箱|email|tel|地址|籍贯|求职意向).*", "", line, flags=re.I).strip()
        first_token = re.split(r"[\s|｜,，;；/]+", candidate)[0].strip()
        if re.fullmatch(r"[\u4e00-\u9fff]{2,4}", first_token) and not detect_section(first_token):
            name = first_token
            break
        if re.fullmatch(r"[\u4e00-\u9fff]{2,4}", candidate) and not detect_section(candidate):
            name = candidate
            break
        if re.fullmatch(r"[A-Za-z][A-Za-z\s]{2,35}", candidate) and len(candidate.split()) <= 4:
            name = candidate
            break
    if not name:
        name = "候选人"
    contact_parts = []
    if phone:
        contact_parts.append(phone.group(0).strip())
    if email:
        contact_parts.append(email.group(0).strip())
    contact_parts.extend(links[:1])
    return {"name": name, "contact": " | ".join(contact_parts), "education": [], "honors": ""}


def parse_education(lines):
    education = []
    current = None
    for line in lines[:12]:
        date, rest = extract_date_from_line(line)
        is_school = has_any(line, ["大学", "学院", "university", "college", "school", "ll.b", "ll.m", "硕士", "本科", "法学", "专业"])
        if date or is_school:
            if current:
                education.append(current)
            parts = re.split(r"\s{2,}|[|｜]", rest or line)
            school = parts[0].strip() if parts else (rest or line)
            major = " ".join(part.strip() for part in parts[1:] if part.strip())
            if not major and len(parts) == 1:
                major = ""
            current = [school, major, date, []]
        elif current:
            current[3].append(line)
    if current:
        education.append(current)
    return [tuple(item) for item in education[:4] if item[0]]


def parse_entries(lines, default_role="相关经历", max_items=8):
    entries = []
    current = None
    for line in lines:
        section = detect_section(line)
        if section:
            continue
        date, rest = extract_date_from_line(line)
        looks_header = bool(date) or (not current and has_any(line, ["公司", "事务所", "银行", "证券", "基金", "大学", "学院", "研究院", "法院", "检察院", "项目", "competition", "university"]))
        if looks_header and len(line) <= 90:
            if current:
                entries.append(current)
            header = rest or line
            parts = [part.strip() for part in re.split(r"\s{2,}|[|｜]", header) if part.strip()]
            org = parts[0] if parts else header
            role = " ".join(parts[1:]) if len(parts) > 1 else default_role
            current = {"org": org[:60], "role": role[:60], "date": date, "raw": [], "bullets": []}
        elif current:
            bullet = clean_bullet_text(line)
            if bullet:
                current["raw"].append(bullet)
        elif line:
            current = {"org": line[:60], "role": default_role, "date": "", "raw": [], "bullets": []}
    if current:
        entries.append(current)
    return entries[:max_items]


def bulletize_external_entry(entry, jd, custom, terms):
    raw = entry.get("raw") or []
    if not raw and entry.get("role"):
        raw = [entry["role"]]
    scored = []
    for idx, text in enumerate(raw):
        scored.append((score_text_for_jd(text, jd, custom, terms), idx, text))
    scored.sort(key=lambda item: (-item[0], item[1]))
    chosen = [text for _, _, text in scored[:3]]
    bullets = []
    for text in chosen:
        if len(text) > 120:
            text = text[:118].rstrip("，；;、 ") + "。"
        if text and not text.endswith(("。", ".", "；", ";")):
            text += "。"
        bullets.append((choose_label(text), text))
    return bullets[:3]


def parse_simple_items(lines, default_title):
    entries = parse_entries(lines, default_role=default_title, max_items=5)
    items = []
    for entry in entries:
        detail = "；".join(entry.get("raw", [])[:2]) or entry.get("role", "")
        items.append((entry.get("org", default_title), entry.get("date", ""), detail[:180]))
    return items


def build_external_resume(jd, selected, resume_text, custom, uploads_text, library, signals, sector, sector_name, terms):
    lines = normalize_resume_lines(resume_text)
    sections = split_resume_sections(lines)
    profile = parse_profile_from_external(lines, sections)
    education = parse_education(sections.get("education", []))
    if education:
        profile["education"] = education
    else:
        profile["education"] = [("教育经历待识别", "", "", ["请检查原简历教育背景是否被成功提取。"])]

    skill_lines = sections.get("skills", [])
    if skill_lines:
        profile["honors"] = "；".join(skill_lines[:2])

    exp_lines = []
    for key in ["experience", "project"]:
        exp_lines.extend(sections.get(key, []))
    entries = parse_entries(exp_lines, default_role="相关经历", max_items=8)
    entries.sort(key=lambda entry: -score_text_for_jd(" ".join([entry.get("org", ""), entry.get("role", ""), *entry.get("raw", [])]), jd, custom, terms))
    experiences = []
    for entry in entries[:5]:
        bullets = bulletize_external_entry(entry, jd, custom, terms)
        if bullets:
            experiences.append({
                "org": entry.get("org", "相关经历"),
                "role": entry.get("role", "相关经历"),
                "date": entry.get("date", ""),
                "bullets": bullets,
            })
    if not experiences:
        fallback_text = compact_text("\n".join(lines[:18]), 260)
        experiences = [{"org": "原简历经历", "role": "待进一步结构化", "date": "", "bullets": [("原始经历摘要：", fallback_text + ("。" if fallback_text else ""))]}]

    competitions = parse_simple_items(sections.get("competition", []), "竞赛经历")[:3]
    research = parse_simple_items(sections.get("research", []), "科研经历")[:3]
    opinions = make_opinions(sector, terms, resume_text, custom, library, {
        "signals": signals,
        "has_external_resume": True,
    })
    opinions.insert(0, {
        "title": "已切换为外部简历事实源",
        "body": "当前预览和导出已基于上传/粘贴的原简历解析生成，不再套用默认模板经历；请检查解析出的姓名、教育和经历标题是否准确。"
    })
    return apply_manual_content_requests({
        "sector": sector,
        "sectorName": sector_name,
        "terms": terms,
        "profile": profile,
        "experiences": experiences,
        "competitions": competitions,
        "research": research,
        "opinions": opinions,
        "library": library,
        "signals": signals,
        "externalResumePending": False,
        "externalParsed": True,
    }, custom)


def detect_sector(jd, selected, signals=None):
    if selected and selected != "auto":
        return selected
    signals = signals or {}
    if signal_active(signals, "金融监管") or signal_active(signals, "金融投资") or signal_active(signals, "绿色投资"):
        return "finance"
    if signal_active(signals, "国央企法务"):
        return "soe"
    normalized = normalize_text(jd)
    scores = {}
    for sector, words in SECTOR_KEYWORDS.items():
        scores[sector] = sum(1 for word in words if word.lower().replace(" ", "") in normalized)
    if not any(scores.values()):
        return "lawfirm"
    return max(scores, key=scores.get)


def extract_terms(jd):
    normalized = normalize_text(jd)
    hits = []
    for label, words in JD_TERMS.items():
        if any(word.lower().replace(" ", "") in normalized for word in words):
            hits.append(label)
    return hits[:8]


def unique_terms(*groups):
    seen = set()
    result = []
    for group in groups:
        for term in group:
            if term not in seen:
                seen.add(term)
                result.append(term)
    return result


def build_tailored_resume(jd, selected="auto", resume_text="", custom="", uploads_text=""):
    library = scan_internship_library(jd, custom)
    jd_custom_signal = "\n".join([jd or "", custom or ""])
    signals = analyze_jd_context(jd or "", custom or "")
    sector = detect_sector(jd_custom_signal, selected, signals)
    jd_terms = extract_terms(jd or "")
    custom_terms = extract_terms(custom or "")
    terms = unique_terms(jd_terms, custom_terms)
    sector_name = {"lawfirm": "律所", "finance": "金融类", "soe": "国央企/法务"}[sector]
    if (resume_text or "").strip():
        return build_external_resume(jd, selected, resume_text, custom, uploads_text, library, signals, sector, sector_name, terms)
    no_data = custom_says_no(custom, ["数据合规", "AI", "AI出海", "技术出口管制", "AIGC"])
    no_finance = custom_says_no(custom, ["金融", "金融监管", "结构性产品", "Final Terms", "绿色私募"])
    no_green_fund = custom_says_no(custom, ["绿色私募", "私募基金", "新能源"])
    no_ldr = custom_says_no(custom, ["LDR", "争议", "诉讼", "仲裁"])
    force_data = custom_says_yes(custom, ["数据合规", "AI", "AI出海", "技术出口管制", "AIGC"])
    force_finreg = custom_says_yes(custom, ["金融监管", "Final Terms", "结构性产品", "equity-linked", "HKEX", "外资银行"])
    force_green_fund = custom_says_yes(custom, ["绿色私募", "私募基金", "新能源", "风能", "地热", "SAF", "金融投资", "投融资"])
    wants_data = not no_data and (force_data or signal_active(signals, "数据合规"))
    wants_finreg = not no_finance and (force_finreg or signal_active(signals, "金融监管"))
    wants_fund = not no_finance and (force_green_fund or signal_active(signals, "金融投资"))
    wants_green = not no_green_fund and (force_green_fund or signal_active(signals, "绿色投资"))
    wants_general_compliance = not no_data and signal_active(signals, "普通合规") and not wants_data and not wants_finreg
    suppressed_terms = set()
    if no_data:
        suppressed_terms.add("数据合规")
    if no_finance:
        suppressed_terms.update(["金融", "基金"])
    if no_ldr:
        suppressed_terms.add("争议")
    terms = [term for term in terms if term not in suppressed_terms]
    has_specific_tailoring = any([
        wants_data,
        wants_finreg,
        wants_fund,
        wants_green,
        wants_general_compliance,
        "知识产权" in terms,
        "争议" in terms,
        "尽调" in terms,
        "英文" in terms,
    ])
    use_sector_bullets = any([
        wants_data,
        wants_finreg,
        wants_fund,
        wants_green,
        wants_general_compliance,
        "知识产权" in terms,
        "争议" in terms,
        "尽调" in terms,
        sector == "soe" and ("合规" in terms or "文书" in terms),
    ])
    reduce_sections = any([
        wants_data,
        wants_finreg,
        wants_fund,
        wants_green,
        "知识产权" in terms,
        "争议" in terms,
        "尽调" in terms,
    ])

    # 面向大众：BASE_PROFILE.experiences 为通用占位（空列表），默认简历不再写死任何个人真实经历。
    # 用户上传简历由 build_external_resume 处理；未上传时由 DeepSeek 基于 JD 生成，或回退到下方通用占位。
    # 历史版本中此处按具体实习公司做 bullet 替换与排序；现仅基于候选人自身上传内容泛化生成，不写死任何机构。
    # bullet 替换与排序，已整体移除，避免把个人简历注入他人定制结果。
    experiences = []

    competition_count = 2 if reduce_sections and sector in {"finance", "soe"} and not wants_data else 3
    competitions = BASE_PROFILE["competitions"][:competition_count]
    research = BASE_PROFILE["research"] if (not reduce_sections or sector != "finance" or wants_data) else BASE_PROFILE["research"][:1]
    if wants_data and not any("模拟立法" in item[0] for item in research):
        research = research + (BASE_PROFILE["research"][1:2] if BASE_PROFILE["research"] else [])

    opinions = make_opinions(sector, terms, resume_text, custom, library, {
        "wants_data": wants_data,
        "wants_finreg": wants_finreg,
        "wants_fund": wants_fund,
        "wants_green": wants_green,
        "force_green_fund": force_green_fund,
        "force_finreg": force_finreg,
        "no_data": no_data,
        "no_finance": no_finance,
        "no_green_fund": no_green_fund,
        "no_ldr": no_ldr,
        "has_specific_tailoring": has_specific_tailoring,
        "use_sector_bullets": use_sector_bullets,
        "wants_general_compliance": wants_general_compliance,
        "signals": signals,
        "has_external_resume": bool((resume_text or "").strip()),
    })
    resume = {
        "sector": sector,
        "sectorName": sector_name,
        "terms": terms,
        "profile": {
            "name": BASE_PROFILE["name"],
            "contact": BASE_PROFILE["contact"],
            "education": BASE_PROFILE["education"],
            "honors": BASE_PROFILE["honors"],
        },
        "experiences": experiences,
        "competitions": competitions,
        "research": research,
        "opinions": opinions,
        "library": library,
        "signals": signals,
        "externalResumePending": bool((resume_text or "").strip()),
    }
    return apply_manual_content_requests(resume, custom)


def make_opinions(sector, terms, resume_text, custom="", library=None, flags=None):
    flags = flags or {}
    signals = flags.get("signals") or {}
    signal_text = signal_summary(signals)
    focus = {
        "lawfirm": "只围绕 JD 明确写到的法律检索、文书、争议、交易、知产或英文能力调整经历；没有出现的方向不主动新增。",
        "finance": "只在 JD 明确出现基金、投融资、结构性产品、金融监管、新能源投资等词时，才前置对应金融经历；否则沿用原简历。",
        "soe": "只在 JD 明确出现合规、监管、公文、合同、内控、政策研究等词时，才改写国央企/法务表达；否则沿用原简历。",
    }[sector]
    delete = {
        "lawfirm": "压缩泛泛奖项和课程，保留 Oralist、核心奖项、法律研究题目。",
        "finance": "删除过细法条和模拟法庭地点，保留能证明研究、英文、商业判断和风险识别的内容。",
        "soe": "删除过细跨境交易术语，保留合规、监管、文书、公共部门和流程优化表述。",
    }[sector]
    gaps = []
    if "英文" in terms:
        gaps.append("JD 明确要求英文，因此保留英文客户沟通、英文竞赛和英文书面表达。")
    if "尽调" in terms:
        gaps.append("JD 明确要求尽调，因此保留交易/债务重组/投资尽调相关经历。")
    if "合规" in terms:
        gaps.append("JD 明确要求合规，因此只选取与监管口径、审批备案、风险边界相关的经历。")
    if not gaps:
        gaps.append("JD 暂未出现需要新增经历的具体关键词，建议保持桌面原始简历，仅做格式和语言微调。")

    targeted = []
    manual = []
    if custom.strip():
        manual.append(f"已读取你的手动要求，并将其作为硬约束处理：{custom.strip()[:160]}")
    if flags.get("has_external_resume"):
        manual.append("已读取上传/粘贴的外部简历；当前本地模式会优先避免使用默认模板内置经历，建议同时上传补充材料并手写明确修改要求。")
    if flags.get("no_data"):
        manual.append("你要求不写/弱化数据合规，因此不会主动加入 AI 出海、技术出口管制或 AIGC 竞赛研究。")
    if flags.get("no_finance"):
        manual.append("你要求不写/弱化金融，因此不会主动加入结构性产品、Final Terms 或私募基金相关经历。")
    if flags.get("no_ldr"):
        manual.append("你要求不写/弱化争议解决，因此相关经历会避开 LDR/doc review 叙述，改写为检索、memo、交易或监管支持。")
    if flags.get("force_finreg"):
        manual.append("你要求突出金融监管，因此 CC 会优先写 Financial Regulation、Final Terms、结构性票据、HKEX 公告和外资银行监管更新。")
    if flags.get("force_green_fund"):
        manual.append("你要求加入绿色私募/新能源，因此会围绕风能、地热、新能源管网、SAF、尽调和投决材料展开，结合你自己的简历补充相关行业研究与投资经历。")

    if flags.get("wants_data"):
        targeted.append("数据合规/AI 出海：增加 AI Lab 技术出口管制、教育类互联网信息服务前置审核、AIGC 模拟立法研究；表述用“检索依据-区分技术/数据类型-归纳审批/备案路径-形成 memo/规则草案”。")
    if flags.get("wants_finreg"):
        targeted.append("金融监管：围绕 structured products、equity-linked notes、autocallable、trigger levels、barrier、redemption amount、Calculation Agent、销售限制及跨境合规写作。")
    if (not flags.get("no_finance")) and (flags.get("wants_fund") or flags.get("wants_green") or "基金" in terms):
        targeted.append("金融投资/基金：围绕风能、地热、新能源管网、SAF 项目写行业研究、投资尽调、投决材料整理；保留债务重组尽调与并购/交易尽调经历，不写 Final Terms。")
    if flags.get("wants_general_compliance"):
        targeted.append("普通合规：不写成数据合规或金融监管；优先呈现 CC 的法规检索、监管口径整理、审批备案/责任边界判断、客户咨询 memo，以及检察院正式文书和公共部门法律审查。")
    if "知识产权" in terms:
        targeted.append("知识产权：保留商标文书、商品包装装潢反不正当竞争、SEP 责任排除和专利文件翻译；弱化检察院经历。")
    if "争议" in terms and not flags.get("no_ldr"):
        targeted.append("争议解决：前置大型跨境投融资争议 doc review、证据目录、核心论点；保留职业放贷、酒店管理合同提前终止、票据/不当得利检索。")
    if sector == "soe" and ("合规" in terms or "文书" in terms):
        targeted.append("国央企/法务：JD 明确出现合规或文书时，才前置检察院正式文书、监管合规和流程规范；未出现时保持原简历排序。")
    if not targeted:
        targeted.append("当前 JD 没有命中需要改写经历的具体词汇，先不新增金融、数据合规、知识产权或绿色私募内容，保持原始简历版本。")

    return [
        {"title": "定位判断", "body": f"当前更像 { {'lawfirm':'律所/法务实习','finance':'金融/投融资相关','soe':'国央企/法务合规'}[sector] } 岗位，命中关键词：{ '、'.join(terms) if terms else '法律检索、文书、尽调' }。"},
        {"title": "语境命中", "body": signal_text or "未命中明确语境；保持原始简历，不主动新增专项经历。"},
        {"title": "手动要求执行", "body": " ".join(manual) if manual else "未填写额外修改要求；当前仅按 JD 关键词和实习资料库进行定制。"},
        {"title": "针对 JD 的实质修改", "body": " ".join(targeted)},
        {"title": "经历排序", "body": focus},
        {"title": "删减方向", "body": delete},
        {"title": "补强提醒", "body": " ".join(gaps)},
        {"title": "实习资料库命中", "body": (library or {}).get("summary") or "暂未命中特定资料；可上传补充文件或在修改要求中写明想强调的项目。"},
    ]


def get_llm_api_base():
    return os.environ.get("LLM_API_BASE", "https://api.deepseek.com").rstrip("/")


def get_deepseek_key():
    return RUNTIME_DEEPSEEK_API_KEY or os.environ.get("LLM_API_KEY", "") or os.environ.get("DEEPSEEK_API_KEY", "")


def get_deepseek_model():
    return RUNTIME_DEEPSEEK_MODEL or os.environ.get("LLM_MODEL", "") or os.environ.get("DEEPSEEK_MODEL", "deepseek-chat")


def deepseek_is_configured():
    return bool(get_deepseek_key())


def set_runtime_deepseek_config(api_key="", model=""):
    global RUNTIME_DEEPSEEK_API_KEY, RUNTIME_DEEPSEEK_MODEL
    api_key = (api_key or "").strip()
    model = (model or "").strip()
    if api_key:
        RUNTIME_DEEPSEEK_API_KEY = api_key
        try:
            LLM_KEY_FILE.write_text(api_key, encoding="utf-8")
        except Exception:
            pass
    if model:
        RUNTIME_DEEPSEEK_MODEL = model
    if not api_key and not model:
        RUNTIME_DEEPSEEK_API_KEY = ""
        RUNTIME_DEEPSEEK_MODEL = ""
        try:
            LLM_KEY_FILE.unlink()
        except Exception:
            pass
    return {"configured": deepseek_is_configured(), "model": "AI", "provider": "AI"}


def resume_for_prompt(resume):
    lines = []
    for exp in resume["experiences"]:
        lines.append(f"{exp['org']} | {exp['role']} | {exp['date']}")
        for label, text in exp["bullets"]:
            lines.append(f"- {label}{text}")
    return "\n".join(lines)


def deepseek_response_text(payload):
    api_key = get_deepseek_key()
    if not api_key:
        raise RuntimeError("未配置 AI 模型 API Key。当前会使用内置英文简历译法。")
    messages = payload.get("messages") or payload.get("input") or []
    request_body = {
        "model": payload.get("model") or get_deepseek_model(),
        "messages": messages,
        "temperature": payload.get("temperature", 0.15),
    }
    if payload.get("response_format"):
        request_body["response_format"] = payload["response_format"]
    last_error = None
    for attempt in range(2):
        req = Request(
            get_llm_api_base() + "/chat/completions",
            data=json.dumps(request_body).encode("utf-8"),
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
            },
            method="POST",
        )
        try:
            with urlopen(req, timeout=60) as response:
                data = json.loads(response.read().decode("utf-8"))
            choices = data.get("choices") or []
            if choices:
                message = choices[0].get("message") or {}
                return (message.get("content") or message.get("reasoning_content") or "").strip()
            return ""
        except HTTPError as exc:
            body = exc.read().decode("utf-8", errors="ignore")
            last_error = body or str(exc)
            if exc.code == 429 and attempt == 0:
                time.sleep(2.0)
                continue
            if exc.code == 429:
                raise RuntimeError("模型请求太频繁或当前 API Key 额度/速率不足。请稍后再试，或换一个有可用额度的 API Key / 更低负载模型。")
            if exc.code == 401:
                raise RuntimeError("DeepSeek API Key 无效或无权限，请重新粘贴正确的 DeepSeek API Key。")
            if exc.code == 402:
                raise RuntimeError("这个 DeepSeek API Key 可能没有可用余额或未开通 API 计费。")
            raise RuntimeError(f"模型调用失败：{last_error[:400]}")
        except URLError as exc:
            last_error = str(exc)
            if attempt == 0:
                time.sleep(1.0)
                continue
            raise RuntimeError(f"网络连接模型失败：{last_error}")
    raise RuntimeError(f"模型调用失败：{last_error or '未知错误'}")


def test_deepseek_connection():
    api_key = get_deepseek_key()
    if not api_key:
        return {"ok": False, "error": "未配置 AI 模型 API Key。请在服务端配置 LLM_API_KEY 后重启。", "model": "AI"}
    probe = {
        "model": get_deepseek_model(),
        "messages": [{"role": "user", "content": "ping"}],
        "max_tokens": 4,
        "temperature": 0,
    }
    req = Request(
        get_llm_api_base() + "/chat/completions",
        data=json.dumps(probe).encode("utf-8"),
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        method="POST",
    )
    try:
        with urlopen(req, timeout=30) as response:
            data = json.loads(response.read().decode("utf-8"))
        if data.get("choices"):
            return {"ok": True, "model": "AI", "provider": "AI"}
        return {"ok": False, "error": "模型返回为空，请重试。", "model": "AI"}
    except HTTPError as exc:
        body = exc.read().decode("utf-8", errors="ignore")
        if exc.code == 401:
            return {"ok": False, "error": "API Key 无效或无权限，请检查 Key。", "model": get_deepseek_model()}
        if exc.code == 402:
            return {"ok": False, "error": "该 Key 无可用余额或未开通计费。", "model": get_deepseek_model()}
        if exc.code == 429:
            return {"ok": False, "error": "请求过于频繁或额度不足，请稍后重试。", "model": get_deepseek_model()}
        return {"ok": False, "error": f"HTTP {exc.code}: {body[:200]}", "model": get_deepseek_model()}
    except URLError as exc:
        return {"ok": False, "error": f"网络连接失败：{exc}", "model": get_deepseek_model()}


def strip_html(text):
    text = re.sub(r"<script[\s\S]*?</script>", " ", text or "", flags=re.I)
    text = re.sub(r"<style[\s\S]*?</style>", " ", text, flags=re.I)
    text = re.sub(r"<[^>]+>", " ", text)
    return compact_text(html_lib.unescape(text), 700)


def clean_result_url(url):
    url = html_lib.unescape(url or "")
    if url.startswith("//"):
        url = "https:" + url
    if "duckduckgo.com/l/" in url or url.startswith("/l/"):
        parsed = urlparse(url)
        params = parse_qs(parsed.query)
        if params.get("uddg"):
            return params["uddg"][0]
    return url


def web_fetch(url, timeout=15):
    req = Request(
        url,
        headers={
            "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X) AppleWebKit/537.36 Chrome/125 Safari/537.36",
            "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
        },
    )
    with urlopen(req, timeout=timeout) as response:
        return response.read().decode("utf-8", errors="ignore")


def search_duckduckgo(query, limit=6):
    url = "https://duckduckgo.com/html/?" + urlencode({"q": query})
    page = web_fetch(url)
    blocks = re.findall(r'<div class="result results_links[^"]*"[\s\S]*?</div>\s*</div>', page)
    results = []
    for block in blocks:
        title_match = re.search(r'<a[^>]+class="result__a"[^>]+href="([^"]+)"[^>]*>([\s\S]*?)</a>', block)
        if not title_match:
            continue
        snippet_match = re.search(r'<a[^>]+class="result__snippet"[^>]*>([\s\S]*?)</a>|<div[^>]+class="result__snippet"[^>]*>([\s\S]*?)</div>', block)
        title = strip_html(title_match.group(2))
        href = clean_result_url(title_match.group(1))
        snippet = strip_html((snippet_match.group(1) or snippet_match.group(2)) if snippet_match else "")
        if title and href:
            results.append({"title": title, "url": href, "snippet": snippet, "source": "DuckDuckGo"})
        if len(results) >= limit:
            break
    return results


def search_bing(query, limit=6):
    rss_url = "https://www.bing.com/search?" + urlencode({"format": "rss", "q": query, "setlang": "zh-CN"})
    rss = web_fetch(rss_url)
    results = []
    try:
        root = etree.fromstring(rss.encode("utf-8"))
        for item in root.xpath(".//item"):
            title = strip_html("".join(item.xpath("./title/text()")))
            href = clean_result_url("".join(item.xpath("./link/text()")))
            snippet = strip_html("".join(item.xpath("./description/text()")))
            if title and href:
                results.append({"title": title, "url": href, "snippet": snippet, "source": "Bing"})
            if len(results) >= limit:
                return results
    except Exception:
        results = []
    url = "https://www.bing.com/search?" + urlencode({"q": query, "setlang": "zh-CN"})
    page = web_fetch(url)
    blocks = re.findall(r'<li class="b_algo"[\s\S]*?</li>', page)
    for block in blocks:
        title_match = re.search(r'<h2[^>]*>\s*<a[^>]+href="([^"]+)"[^>]*>([\s\S]*?)</a>', block)
        snippet_match = re.search(r"<p[^>]*>([\s\S]*?)</p>", block)
        if not title_match:
            continue
        title = strip_html(title_match.group(2))
        href = clean_result_url(title_match.group(1))
        snippet = strip_html(snippet_match.group(1) if snippet_match else "")
        if title and href:
            results.append({"title": title, "url": href, "snippet": snippet, "source": "Bing"})
        if len(results) >= limit:
            break
    return results


def make_web_query(jd, custom):
    signals = analyze_jd_context(jd or "", custom or "")
    priority = []
    for item in signals.values():
        priority.extend(item.get("jdHits") or [])
        priority.extend(item.get("customHits") or [])
    priority.extend(query_terms_from_text(jd, custom))
    cleaned = []
    for term in priority:
        term = re.sub(r"\s+", " ", str(term)).strip()
        if not term or normalize_text(term) in QUERY_STOPWORDS:
            continue
        if len(term) > 18 and re.fullmatch(r"[\u4e00-\u9fff]+", term):
            continue
        if term.lower() not in [item.lower() for item in cleaned]:
            cleaned.append(term)
    if not cleaned:
        return "法务 合规 金融 律所 法律实务 监管"
    return " ".join(cleaned[:8] + ["法律实务", "监管"])


def make_web_queries(jd, custom):
    signals = analyze_jd_context(jd or "", custom or "")
    queries = []
    if signal_active(signals, "数据合规") or custom_mentions(custom, ["模拟立法", "生成式人工智能", "AI出海", "技术出口管制"]):
        queries.extend([
            "生成式人工智能服务管理暂行办法 技术出口管制 个人信息保护 数据出境 法律实务",
            "site:cac.gov.cn 生成式人工智能 服务管理 暂行办法",
            "个人信息保护法 数据出境 个人信息保护 法律合规 实务",
        ])
    if signal_active(signals, "金融监管"):
        queries.extend([
            "structured products final terms equity-linked notes selling restrictions legal",
            "HKEX structured products foreign banks regulatory announcements final terms",
            "autocallable notes final terms calculation agent selling restrictions",
        ])
    if signal_active(signals, "金融投资") or signal_active(signals, "绿色投资"):
        queries.extend([
            "私募基金 尽职调查 投资决策 新能源 风能 地热 法律实务",
            "绿色投资 新能源 项目尽职调查 风险评估 投决 法务",
        ])
    if signal_active(signals, "知识产权"):
        queries.append("商标 商品包装装潢 反不正当竞争 知识产权 法律实务")
    if signal_active(signals, "争议解决"):
        queries.append("投融资争议 股权回售权 put option doc review 证据 法律实务")
    queries.append(make_web_query(jd, custom))
    unique = []
    seen = set()
    for query in queries:
        key = normalize_text(query)
        if key not in seen:
            seen.add(key)
            unique.append(query)
    return unique[:5]


def useful_web_term(term):
    key = normalize_text(term)
    if not key or key in WEB_WEAK_TERMS or key in QUERY_STOPWORDS or re.fullmatch(r"20\d{2}|\d+", key):
        return False
    if key in WEB_SHORT_ALLOWED:
        return True
    if re.search(r"[a-z]", key):
        return len(key) >= 5
    return len(key) >= 4


def web_result_score(item, terms):
    haystack = normalize_text(f"{item.get('title', '')} {item.get('snippet', '')}")
    score = sum(1 for term in terms if normalize_text(term) in haystack)
    core_phrases = [
        "生成式人工智能", "技术出口管制", "个人信息保护", "个人信息保护法", "数据出境", "网络安全法", "数据安全法",
        "finalterms", "structuredproducts", "equitylinked", "autocallable", "hkex", "calculationagent",
        "私募基金", "尽职调查", "新能源", "风能", "地热", "投资决策", "投决", "风险评估",
        "商标", "专利", "不正当竞争", "股权回售", "putoption",
    ]
    core_hits = sum(1 for phrase in core_phrases if phrase in haystack)
    domain = urlparse(item.get("url", "")).netloc.lower()
    authoritative_domains = ("gov.cn", "cac.gov.cn", "miit.gov.cn", "moj.gov.cn", "mps.gov.cn", "hkex.com.hk", "sfc.hk", "pbc.gov.cn", "csrc.gov.cn")
    if score and core_hits:
        score += core_hits * 2
    elif score < 2:
        return 0
    if score and any(domain.endswith(d) or d in domain for d in authoritative_domains):
        score += 2
    return score


def run_web_research(jd, custom):
    queries = make_web_queries(jd, custom)
    errors = []
    results = []
    terms = [term for term in query_terms_from_text(jd, custom, *queries) if useful_web_term(term)]
    for query in queries:
        for searcher in (search_bing, search_duckduckgo):
            try:
                found = searcher(query)
                for item in found:
                    scored = {**item, "query": query, "score": web_result_score(item, terms)}
                    if scored["score"] and item.get("url") not in {existing.get("url") for existing in results}:
                        results.append(scored)
                if len(results) >= 6:
                    break
            except Exception as exc:
                errors.append(str(exc))
        if len(results) >= 6:
            break
    results.sort(key=lambda item: (-item["score"], item["title"]))
    results = results[:6]
    summary = "；".join(f"{item['title']}：{compact_text(item.get('snippet', ''), 180)}" for item in results[:4] if item.get("snippet"))
    summary = compact_text(summary, 900)
    if not summary and results:
        summary = "；".join(item["title"] for item in results[:4])
    return {
        "query": "；".join(queries),
        "count": len(results),
        "results": results,
        "summary": summary or "联网检索没有返回可用摘要。可以换一个更具体的 JD 或在修改要求里写公司名/业务方向。",
        "errors": errors[:2],
    }


def web_research_for_prompt(research):
    lines = [f"检索式：{research.get('query', '')}"]
    for item in research.get("results", [])[:6]:
        lines.append(f"- {item.get('title')} | {item.get('url')} | {item.get('snippet')}")
    return "\n".join(lines)


def call_deepseek_advice(jd, custom, resume):
    model = get_deepseek_model()
    system = (
        "你是中国法学生求职简历定制助手。必须严格根据JD明示关键词和用户修改要求改简历；"
        "如果JD没有出现某类关键词，不要主动新增对应经历。不得虚构经历，不得夸大。"
        "输出中文，给出3-6条具体修改意见，并说明哪些经历应保留、删除、前置或改写。"
    )
    user = (
        f"JD：\n{jd or '未提供'}\n\n"
        f"用户手动修改要求：\n{custom or '未提供'}\n\n"
        f"当前候选简历经历：\n{resume_for_prompt(resume)}\n\n"
        "请判断是否需要改写。特别注意：只有JD或用户要求明确提到的数据合规、金融监管、Final Terms、私募基金、知识产权、争议解决等词，才可调用对应经历。"
    )
    payload = {
        "model": model,
        "input": [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        "temperature": 0.2,
    }
    return deepseek_response_text(payload) or "DeepSeek 已返回结果，但未解析出文字内容。"


def extract_json_object(text):
    text = (text or "").strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```$", "", text)
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        start = text.find("{")
        end = text.rfind("}")
        if start != -1 and end != -1 and end > start:
            return json.loads(text[start : end + 1])
        raise


def normalize_bullets(bullets):
    normalized = []
    for item in bullets or []:
        if isinstance(item, (list, tuple)) and len(item) >= 2:
            label = str(item[0]).strip()
            text = str(item[1]).strip()
        elif isinstance(item, dict):
            label = str(item.get("label", "")).strip()
            text = str(item.get("text", "")).strip()
        else:
            continue
        if label and not label.endswith("："):
            label += "："
        if label and text:
            normalized.append((label, text))
    return normalized


def normalize_profile(profile, base_profile):
    profile = profile if isinstance(profile, dict) else {}
    base_profile = base_profile if isinstance(base_profile, dict) else {}
    normalized = {
        "name": str(profile.get("name") or base_profile.get("name") or "").strip(),
        "contact": str(profile.get("contact") or base_profile.get("contact") or "").strip(),
        "honors": str(profile.get("honors") or base_profile.get("honors") or "").strip(),
        "education": [],
    }
    education = profile.get("education") if isinstance(profile.get("education"), list) else base_profile.get("education", [])
    for item in education or []:
        if isinstance(item, dict):
            school = str(item.get("school", "")).strip()
            major = str(item.get("major", "")).strip()
            date = str(item.get("date", "")).strip()
            details = item.get("details", [])
        elif isinstance(item, (list, tuple)) and len(item) >= 4:
            school, major, date, details = item[0], item[1], item[2], item[3]
        elif isinstance(item, (list, tuple)) and len(item) >= 3:
            school, major, date, details = item[0], item[1], item[2], []
        else:
            continue
        if isinstance(details, str):
            details = [details]
        details = [str(detail).strip() for detail in (details or []) if str(detail).strip()]
        if str(school).strip() or str(major).strip():
            normalized["education"].append((str(school).strip(), str(major).strip(), str(date).strip(), details))
    if not normalized["education"]:
        normalized["education"] = base_profile.get("education", [])
    return normalized


def normalize_resume(candidate, base):
    candidate = candidate or {}
    if "resume" in candidate and isinstance(candidate["resume"], dict):
        candidate = candidate["resume"]
    resume = {**base}
    for key in ["sector", "sectorName", "terms", "library", "signals", "webResearch"]:
        if key in candidate:
            resume[key] = candidate[key]
    if "profile" in candidate:
        resume["profile"] = normalize_profile(candidate.get("profile"), base.get("profile", {}))
    experiences = []
    for exp in candidate.get("experiences", []) if isinstance(candidate.get("experiences"), list) else []:
        if not isinstance(exp, dict):
            continue
        bullets = normalize_bullets(exp.get("bullets"))
        if not bullets:
            continue
        experiences.append({
            "org": str(exp.get("org", "")).strip(),
            "role": str(exp.get("role", "")).strip(),
            "date": str(exp.get("date", "")).strip(),
            "bullets": bullets[:3],
        })
    if experiences:
        resume["experiences"] = experiences[:5]
    for key in ["competitions", "research"]:
        items = []
        for item in candidate.get(key, []) if isinstance(candidate.get(key), list) else []:
            if isinstance(item, (list, tuple)) and len(item) >= 3:
                items.append((str(item[0]).strip(), str(item[1]).strip(), str(item[2]).strip()))
            elif isinstance(item, dict):
                items.append((str(item.get("title", "")).strip(), str(item.get("date", "")).strip(), str(item.get("detail", "")).strip()))
        if items:
            resume[key] = items[:3]
    opinions = candidate.get("opinions")
    if isinstance(opinions, list) and opinions:
        normalized_opinions = []
        for item in opinions:
            if isinstance(item, dict):
                normalized_opinions.append({"title": str(item.get("title", "改写说明")), "body": str(item.get("body", ""))})
        if normalized_opinions:
            resume["opinions"] = normalized_opinions + [op for op in base.get("opinions", []) if op.get("title") == "语境命中"]
    return resume


def call_deepseek_rewrite_resume(jd, custom, base_resume, source_resume_text="", uploads_text=""):
    model = get_deepseek_model()
    allowed_labels = sorted(set(EN_BULLETS.keys()) | {
        "行业研究：", "投资尽调：", "合规研究：", "法律检索：", "文书写作：", "交易支持：",
        "风险识别：", "客户沟通：", "英文写作：", "项目材料整理：",
    })
    system = (
        "你是一个严谨的中文法律求职简历改写器。你必须直接生成一份可进入Word的一页简历JSON，"
        "不是只给建议。严格根据JD和用户修改要求现场选择经历和重写bullet。"
        "用户修改要求的优先级高于你的默认判断；只要不违背事实和不构成虚构，就必须执行。"
        "如果用户要求丰富某段经历、删除某段经历、前置某段经历、增加某项研究或改写某个方向，必须反映到输出resume正文和opinions里。"
        "不得虚构经历、客户名称、项目名称或数据；可以在已提供经历范围内重组、删减、改写表达。"
        "如果用户提供了原始简历文本，必须以该原始简历为唯一事实来源，不得使用默认候选简历中的人物、学校、机构、经历。"
        "如果JD只是泛泛岗位大类，不要套刻板印象；如果JD具体写到合规、金融投资、金融监管、数据合规、知产或争议，才调用对应素材。"
        "bullet要用STAR逻辑，但不要写Situation/Task字样；每条要具体、专业、可落地。"
        "必须在输出的 resume.profile 中填写 enName 字段，为候选人姓名的英文拼音（姓在前、空格分隔、每个词首字母大写，例如 张三→Zhang San、李四→Li Si）；即使原简历是中文，也必须生成英文名，不得留空或返回中文。"
        "输出必须是JSON对象，不要Markdown。"
    )
    user = {
        "JD": jd or "",
        "修改要求": custom or "",
        "必须执行修改要求": bool((custom or "").strip()),
        "用户上传或粘贴的原始简历文本": source_resume_text[:12000] if source_resume_text else "",
        "用户上传的补充材料文本": uploads_text[:12000] if uploads_text else "",
        "默认候选简历": base_resume,
        "桌面资料命中": base_resume.get("library", {}),
        "联网检索结果": base_resume.get("webResearch", {}),
        "事实来源规则": "如果“用户上传或粘贴的原始简历文本”非空，只能使用该文本和用户上传补充材料中的事实；默认候选简历只能作为格式参考，禁止混入不属于该候选人的人物、学校、机构、经历。如果原始简历文本为空，才使用默认候选简历事实。",
        "检索材料使用规则": "桌面资料和联网检索结果只能用于理解JD、行业语境和措辞；除非材料明确来自该候选人的经历，否则不得据此捏造候选人没有做过的经历。",
        "输出质量要求": "opinions中必须单列一条“手动修改要求如何执行”，说明修改要求具体改进到了哪些经历、竞赛或科研条目；resume必须体现这些改动，不能只在意见里说明。",
        "可用标签示例": allowed_labels,
        "输出JSON格式": {
            "summary": "一句话说明改写策略",
            "resume": {
                "sector": base_resume.get("sector"),
                "sectorName": base_resume.get("sectorName"),
                "terms": base_resume.get("terms", []),
                "profile": {
                    "name": "候选人姓名",
                    "enName": "Li Si",
                    "contact": "电话 | 邮箱 | 链接",
                    "education": [{"school": "学校", "major": "专业/学位", "date": "时间", "details": ["成绩/荣誉/课程"]}],
                    "honors": "荣誉奖项",
                },
                "experiences": [
                    {"org": "机构名", "role": "岗位名", "date": "时间", "bullets": [["标签：", "正文"]]}
                ],
                "competitions": [["标题", "时间", "内容"]],
                "research": [["标题", "时间", "内容"]],
                "opinions": [{"title": "DeepSeek改写说明", "body": "说明如何根据JD和修改要求改写"}],
            },
        },
    }
    payload = {
        "model": model,
        "input": [
            {"role": "system", "content": system},
            {"role": "user", "content": json.dumps(user, ensure_ascii=False)},
        ],
        "temperature": 0.35,
        "response_format": {"type": "json_object"},
    }
    text = deepseek_response_text(payload)
    data = extract_json_object(text)
    rewritten = normalize_resume(data, base_resume)
    if isinstance(data, dict) and isinstance(data.get("resume"), dict):
        en = data["resume"].get("profile", {}).get("enName")
        if en:
            rewritten.setdefault("profile", {})["enName"] = en
    summary = data.get("summary") if isinstance(data, dict) else ""
    return {"answer": summary or "DeepSeek 已将修改要求写入简历正文。", "resume": rewritten}


def resume_from_payload_or_build(payload, custom, uploads_text):
    override = payload.get("resumeOverride")
    if isinstance(override, dict) and override.get("experiences"):
        return normalize_resume(override, build_tailored_resume(payload.get("jd", ""), payload.get("sector", "auto"), payload.get("resumeText", ""), custom, uploads_text))
    if (payload.get("resumeText") or "").strip():
        raise RuntimeError("已上传/粘贴外部简历。为了避免生成错人的简历，请先点击“生成修改意见”或“联网检索并改写”，确认预览内容已经基于这份外部简历后再导出 Word。")
    return build_tailored_resume(payload.get("jd", ""), payload.get("sector", "auto"), payload.get("resumeText", ""), custom, uploads_text)


def build_or_deepseek_rewrite(payload, custom, uploads_text, web_research=None):
    jd = payload.get("jd", "")
    sector = payload.get("sector", "auto")
    resume_text = payload.get("resumeText", "")
    base_resume = build_tailored_resume(jd, sector, resume_text, custom, uploads_text)
    if web_research:
        base_resume["webResearch"] = web_research
        base_resume["opinions"].insert(0, {"title": "联网检索命中", "body": web_research.get("summary", "")})
    if not deepseek_is_configured():
        base_resume["mode"] = "local"
        return {"resume": base_resume, "answer": "未配置 AI 模型 Key，已使用本地规则生成。", "mode": "local", "safeToExport": True}
    try:
        result = call_deepseek_rewrite_resume(jd, custom, base_resume, resume_text, uploads_text)
        rewritten = result["resume"]
        rewritten["mode"] = "model"
        rewritten["opinions"].insert(0, {
            "title": "DeepSeek已接入",
            "body": "已根据 JD、原简历、上传材料和手动修改意见生成可直接替换进简历的内容。",
        })
        return {"resume": rewritten, "answer": result.get("answer") or "AI 已完成定制改写。", "mode": "model", "safeToExport": True}
    except Exception as exc:
        base_resume["mode"] = "local"
        base_resume["opinions"].insert(0, {
            "title": "DeepSeek调用失败，已兜底",
            "body": f"{str(exc)} 当前已使用本地规则生成，保存正确 Key 后可重新生成。",
        })
        return {"resume": base_resume, "answer": f"AI 调用失败，已使用本地规则兜底：{exc}", "mode": "local", "safeToExport": True}


# ---- 统一的 Word 版式（中英文共用）----
ACCENT = RGBColor(0x1F, 0x3A, 0x5F)        # 深藏青，专业稳重
CONTACT_RGB = RGBColor(0x59, 0x59, 0x59)   # 联系方式用灰
ZH_HEAD_FONT = "黑体"                        # 章节标题 / 姓名
ZH_BODY_FONT = "宋体"                        # 正文
DOC_TOP = Cm(0.75)
DOC_BOTTOM = Cm(0.7)
DOC_LEFT = Cm(1.05)
DOC_RIGHT = Cm(1.05)


def font_run(run, size=9, bold=False, font=None, latin="Times New Roman"):
    cjk = font or ZH_BODY_FONT
    run.font.name = cjk
    run._element.rPr.rFonts.set(qn("w:eastAsia"), cjk)
    run._element.rPr.rFonts.set(qn("w:ascii"), latin)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), latin)
    run.font.size = Pt(size)
    run.font.bold = bold
    return run


def setup_doc_page(doc, cjk_font, latin_font):
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = DOC_TOP
    sec.bottom_margin = DOC_BOTTOM
    sec.left_margin = DOC_LEFT
    sec.right_margin = DOC_RIGHT
    for style_name in ["Normal", "List Bullet"]:
        style = doc.styles[style_name]
        style.font.name = cjk_font
        style._element.rPr.rFonts.set(qn("w:eastAsia"), cjk_font)
        style._element.rPr.rFonts.set(qn("w:ascii"), latin_font)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), latin_font)
        style.font.size = Pt(9)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 0.95


def add_name(doc, name, font=ZH_HEAD_FONT, size=16):
    p = doc.add_paragraph()
    para_fmt(p, after=0.12, align=WD_ALIGN_PARAGRAPH.CENTER, line=0.9)
    r = font_run(p.add_run(name), size=size, bold=True, font=font, latin="Times New Roman")
    r.font.color.rgb = ACCENT
    bottom_border(p, color="1F3A5F")


def add_contact(doc, contact, font=ZH_BODY_FONT, size=9):
    p = doc.add_paragraph()
    para_fmt(p, after=0.35, align=WD_ALIGN_PARAGRAPH.CENTER, line=0.9)
    r = font_run(p.add_run(contact), size=size, font=font, latin="Times New Roman")
    r.font.color.rgb = CONTACT_RGB


def section_title(doc, title, font=ZH_HEAD_FONT, size=11.5):
    p = doc.add_paragraph()
    para_fmt(p, before=1.2, after=0.25, line=0.9)
    r = font_run(p.add_run(title), size=size, bold=True, font=font, latin="Times New Roman")
    r.font.color.rgb = ACCENT
    bottom_border(p, color="1F3A5F")


def para_fmt(p, before=0, after=0, line=0.86, align=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None:
        p.alignment = align


def bottom_border(p, color="000000"):
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)


def add_section(doc, title):
    section_title(doc, title)


def add_line(doc, left, center="", right="", size=8.9, bold=True):
    p = doc.add_paragraph()
    para_fmt(p, before=0.2, after=0, line=0.9)
    tabs = p.paragraph_format.tab_stops
    tabs.add_tab_stop(Cm(9.15), WD_ALIGN_PARAGRAPH.CENTER)
    tabs.add_tab_stop(Cm(18.35), WD_ALIGN_PARAGRAPH.RIGHT)
    font_run(p.add_run(left), size=size, bold=bold)
    if center:
        font_run(p.add_run("\t" + center), size=size)
    if right:
        font_run(p.add_run(("\t" if center else "\t\t") + right), size=size)


def add_plain(doc, text, size=8.1, indent=0.32):
    if not text:
        return
    p = doc.add_paragraph()
    para_fmt(p, after=0, line=0.92)
    p.paragraph_format.left_indent = Cm(indent)
    font_run(p.add_run(text), size=size)


def add_bullet(doc, label, text, size=8.0):
    p = doc.add_paragraph(style="List Bullet")
    para_fmt(p, after=0, line=0.9)
    p.paragraph_format.left_indent = Cm(0.58)
    p.paragraph_format.first_line_indent = Cm(-0.2)
    font_run(p.add_run(label), size=size, bold=True)
    font_run(p.add_run(text), size=size)


def build_docx(resume):
    doc = Document()
    setup_doc_page(doc, ZH_BODY_FONT, "Times New Roman")

    bullet_count = sum(len(exp.get("bullets", [])) for exp in resume.get("experiences", []))
    dense = bullet_count + len(resume.get("competitions", [])) + len(resume.get("research", [])) > 14
    body_size = 8.6 if dense else 9.0
    line_size = 9.2 if dense else 9.5
    detail_size = 8.4 if dense else 8.7

    add_name(doc, resume["profile"]["name"], font=ZH_HEAD_FONT, size=16)
    add_contact(doc, resume["profile"].get("contact", ""), font=ZH_BODY_FONT, size=9)

    section_title(doc, "教育背景")
    for school, major, date, details in resume["profile"].get("education", []):
        add_line(doc, school, major, date, size=line_size)
        for detail in details:
            add_plain(doc, detail, size=detail_size)
    add_plain(doc, resume["profile"].get("honors", ""), size=detail_size)

    section_title(doc, "实习经历")
    for exp in resume.get("experiences", []):
        add_line(doc, f'{exp.get("org", "")} | {exp.get("role", "")}', right=exp.get("date", ""), size=line_size)
        for label, text in exp.get("bullets", []):
            add_bullet(doc, label, text, size=body_size)

    if resume.get("competitions"):
        section_title(doc, "竞赛经历")
        for title, date, detail in resume.get("competitions", []):
            add_line(doc, title, right=date, size=line_size)
            add_plain(doc, detail, size=detail_size)

    if resume.get("research"):
        section_title(doc, "科研经历")
        for title, date, detail in resume.get("research", []):
            add_line(doc, title, right=date, size=line_size)
            add_plain(doc, detail, size=detail_size)

    section_title(doc, "技能与兴趣")
    if resume.get("externalParsed"):
        add_plain(doc, resume["profile"].get("honors", "") or "技能：请根据原简历补充语言、证书、软件等信息。", size=detail_size)
    else:
        add_plain(doc, "技能与兴趣：请上传你的简历文件，或在“补充材料”中填写技能、语言、证书与兴趣，以生成个性化内容。", size=detail_size)

    filename = f"JD定制简历_{time.strftime('%Y%m%d_%H%M%S')}.docx"
    output_path = OUTPUTS / filename
    doc.save(output_path)
    return output_path


def wqn(local):
    return f"{{{W_NS}}}{local}"


def load_doc_root(path):
    with ZipFile(path) as z:
        return etree.fromstring(z.read("word/document.xml"))


def first_ooxml_rpr(p):
    for r in p.xpath("./w:r", namespaces=NS):
        rpr = r.find("w:rPr", namespaces=NS)
        if rpr is not None:
            return etree.fromstring(etree.tostring(rpr))
    return etree.Element(wqn("rPr"))


def ooxml_set_text(p, text):
    rpr = first_ooxml_rpr(p)
    for r in p.xpath("./w:r", namespaces=NS):
        p.remove(r)
    r = etree.SubElement(p, wqn("r"))
    r.append(rpr)
    pieces = text.split("\t")
    for idx, piece in enumerate(pieces):
        if piece:
            t = etree.SubElement(r, wqn("t"))
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t.text = piece
        if idx < len(pieces) - 1:
            etree.SubElement(r, wqn("tab"))


def ooxml_clone(paras, idx, text):
    p = etree.fromstring(etree.tostring(paras[idx]))
    ooxml_set_text(p, text)
    return p


def write_ooxml_docx(template, out, paragraphs):
    with ZipFile(template, "r") as zin:
        files = {name: zin.read(name) for name in zin.namelist()}
    root = etree.fromstring(files["word/document.xml"])
    body = root.find("w:body", namespaces=NS)
    sect_pr = body.find("w:sectPr", namespaces=NS)
    for child in list(body):
        body.remove(child)
    for p in paragraphs:
        body.append(p)
    if sect_pr is not None:
        body.append(sect_pr)
    files["word/document.xml"] = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
    with ZipFile(out, "w", ZIP_DEFLATED) as zout:
        for name, data in files.items():
            zout.writestr(name, data)
    return out


def build_docx_template_zh(resume):
    return build_docx(resume)

def build_docx_template_en(resume):
    return build_docx_en(resume)

def en_font_run(run, size=9.5, bold=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    return run


def en_section(doc, title):
    p = doc.add_paragraph()
    para_fmt(p, before=1.1, after=0.25, line=0.9)
    en_font_run(p.add_run(title), size=11.5, bold=True, color=ACCENT)
    bottom_border(p, color="1F3A5F")


def en_line(doc, left, right="", size=9.5, bold=True):
    p = doc.add_paragraph()
    para_fmt(p, before=0, after=0, line=0.82)
    tabs = p.paragraph_format.tab_stops
    tabs.add_tab_stop(Cm(18.05), WD_ALIGN_PARAGRAPH.RIGHT)
    en_font_run(p.add_run(left), size=size, bold=bold)
    if right:
        en_font_run(p.add_run("\t" + right), size=size, bold=False)


def en_detail(doc, text, size=9.2):
    p = doc.add_paragraph()
    para_fmt(p, before=0, after=0, line=0.80)
    p.paragraph_format.left_indent = Cm(0.35)
    en_font_run(p.add_run(text), size=size)


def en_bullet(doc, label, text, size=9.1):
    p = doc.add_paragraph(style="List Bullet")
    para_fmt(p, before=0, after=0, line=0.78)
    p.paragraph_format.left_indent = Cm(0.55)
    p.paragraph_format.first_line_indent = Cm(-0.2)
    en_font_run(p.add_run(label), size=size, bold=True)
    en_font_run(p.add_run(text), size=size)


# 面向大众：英文机构映射只保留通用规则/占位，不再写死任何个人真实实习机构
# （具体实习机构名单已移除），避免把个人经历翻译成他人简历。
EN_TRANSLATIONS = {}

EN_SCHOOLS = {
    "北京大学": "Peking University",
    "清华大学": "Tsinghua University",
    "中国人民大学": "Renmin University of China",
    "中国政法大学": "China University of Political Science and Law",
    "复旦大学": "Fudan University",
    "上海交通大学": "Shanghai Jiao Tong University",
    "武汉大学": "Wuhan University",
    "华东政法大学": "East China University of Political Science and Law",
    "对外经济贸易大学": "University of International Business and Economics",
    "南京大学": "Nanjing University",
    "浙江大学": "Zhejiang University",
    "意大利博洛尼亚大学，公派交流项目": "Alma Mater Studiorum - University of Bologna | Exchange Program",
}

EN_ROLE_TERMS = [
    ("数据合规实习生", "Data Compliance Intern"),
    ("法务实习生", "Legal Intern"),
    ("法律实习生", "Legal Intern"),
    ("律师助理", "Legal Assistant"),
    ("实习生", "Intern"),
    ("项目核心成员", "Core Member"),
    ("负责人", "Lead Member"),
    ("成员", "Member"),
]

EN_ORG_TERMS = [
    ("律师事务所", "Law Firm"),
    ("检察院", "Procuratorate"),
    ("法院", "Court"),
    ("银行", "Bank"),
    ("证券", "Securities"),
    ("基金", "Fund"),
    ("投资", "Investment"),
    ("公司", "Company"),
    ("大学", "University"),
    ("学院", "School"),
]

EN_BULLET_LABELS = {
    "数据与技术合规：": "Data and technology compliance: ",
    "金融与交易支持：": "Finance and transaction support: ",
    "合规与风险识别：": "Compliance and risk identification: ",
    "研究分析：": "Research and analysis: ",
    "文书写作：": "Legal drafting: ",
    "英文与跨境支持：": "English and cross-border support: ",
    "工作内容：": "Work scope: ",
    "原始经历摘要：": "Experience summary: ",
}


def has_chinese(text):
    return bool(re.search(r"[\u4e00-\u9fff]", text or ""))


def en_name(name):
    if not has_chinese(name):
        return name or "Candidate"
    return name or "Candidate"


def en_replace_terms(text, pairs):
    result = text or ""
    for zh, en in pairs:
        result = result.replace(zh, en)
    return result


def en_org(text):
    if text in EN_TRANSLATIONS:
        return EN_TRANSLATIONS[text][0]
    if not has_chinese(text):
        return text or "Organization"
    for zh, en in EN_SCHOOLS.items():
        text = text.replace(zh, en)
    text = en_replace_terms(text, EN_ORG_TERMS)
    return text or "Organization"


def en_role(text):
    if not has_chinese(text):
        return text or "Intern"
    return en_replace_terms(text, EN_ROLE_TERMS) or "Intern"


def en_school_major(school, major):
    school_en = EN_SCHOOLS.get(school, en_org(school))
    major_text = major or ""
    if not major_text:
        return school_en
    replacements = [
        ("民商法学硕", "LL.M. in Civil and Commercial Law"),
        ("民商法", "Civil and Commercial Law"),
        ("法学本科", "LL.B. / Law"),
        ("法学", "Law"),
        ("法律", "Law"),
        ("硕士", "Master's Program"),
        ("本科", "Bachelor's Program"),
        ("专业", ""),
    ]
    major_en = en_replace_terms(major_text, replacements)
    return f"{school_en} | {major_en}"


def en_detail_text(text):
    if not text:
        return ""
    if not has_chinese(text):
        return text.replace("：", ": ")
    replacements = [
        ("2024年通过国家司法考试", "Passed PRC Legal Professional Qualification Examination (2024)"),
        ("GPA：", "GPA: "),
        ("综合排名", "Comprehensive Ranking"),
        ("专业排名", "Major Ranking"),
        ("专必均分", "Major Required Courses Average"),
        ("雅思", "IELTS "),
        ("通过国家司法考试", "Passed PRC Legal Professional Qualification Examination"),
        ("荣誉奖项", "Honors"),
        ("校级", "University-level "),
        ("竞赛一等奖学金", "First Prize Competition Scholarship"),
        ("连续两年学业奖学金", "Academic Scholarship for Two Consecutive Years"),
        ("帕特森基金会奖学金", "Patterson Foundation Scholarship"),
        ("优秀毕业生", "Outstanding Graduate"),
        ("三好学生", "Outstanding Student"),
        ("奖学金", "Scholarship"),
        ("法律检索", "Legal Research"),
        ("案例检索", "Case Research"),
        ("普通话", "Mandarin"),
        ("母语", "Native"),
    ]
    text = en_replace_terms(text, replacements)
    text = text.replace("：", ": ").replace("；", "; ").replace("，", ", ").replace("、", "; ")
    return re.sub(r"\s+", " ", text).strip()


def en_bullet_text(label, text):
    if label in EN_BULLETS:
        return EN_BULLETS[label]
    if not has_chinese(text):
        return EN_BULLET_LABELS.get(label, label.replace("：", ": ")), text
    normalized = normalize_text(f"{label}{text}")
    parts = []
    if has_any(normalized, ["个人信息", "数据出境", "数据合规", "隐私", "网络安全"]):
        parts.append("researched personal information protection, data export compliance and cybersecurity requirements")
    if has_any(normalized, ["生成式人工智能", "aigc", "ai", "算法", "技术出口管制", "出口管制"]):
        parts.append("analyzed AI-related regulatory issues and technology export-control requirements")
    if has_any(normalized, ["法律检索", "法规", "监管", "政策", "案例", "司法解释"]):
        parts.append("conducted legal and regulatory research and summarized applicable rules")
    if has_any(normalized, ["合同", "条款", "审查", "审核"]):
        parts.append("reviewed contract clauses and identified legal risk points")
    if has_any(normalized, ["尽调", "工商", "股权", "涉诉", "负面信息", "知识产权权属"]):
        parts.append("organized due diligence materials on corporate registration, shareholding, litigation and IP issues")
    if has_any(normalized, ["备忘录", "memo", "报告", "文书", "起草", "撰写", "邮件"]):
        parts.append("prepared memoranda, reports and drafting materials for lawyer or client review")
    if has_any(normalized, ["金融", "投资", "基金", "投融资", "估值", "行业研究", "投决"]):
        parts.append("supported financial, investment and industry research work")
    if has_any(normalized, ["争议", "诉讼", "仲裁", "证据", "庭审"]):
        parts.append("reviewed dispute materials and organized evidence and procedural issues")
    if has_any(normalized, ["商标", "专利", "著作权", "知识产权", "侵权"]):
        parts.append("analyzed intellectual property issues including trademarks, patents and infringement risks")
    if has_any(normalized, ["英文", "英语", "翻译", "跨境", "境外"]):
        parts.append("translated materials and supported English or cross-border communications")
    if not parts:
        parts.append("organized project materials, summarized key issues and supported legal research or drafting work")
    sentence = "; ".join(parts)
    sentence = sentence[0].upper() + sentence[1:] + "."
    return EN_BULLET_LABELS.get(label, "Work scope: "), sentence


def en_item(title, detail="", kind="item"):
    if title in EN_COMPETITIONS:
        return EN_COMPETITIONS[title]
    if title in EN_RESEARCH:
        return EN_RESEARCH[title]
    if not has_chinese(f"{title}{detail}"):
        return title, detail
    title_en = en_replace_terms(title, [
        ("竞赛经历", "Competition Experience"),
        ("科研经历", "Research Experience"),
        ("项目经历", "Project Experience"),
        ("全国法科学生模拟立法大赛", "National Legislative Simulation Competition"),
        ("模拟立法", "Legislative Simulation"),
        ("模拟法庭", "Moot Court"),
        ("项目核心成员", "Core Member"),
        ("独著", "Sole Author"),
        ("论文", "Research Paper"),
        ("研究", "Research"),
    ])
    _, detail_en = en_bullet_text("研究分析：", detail or title)
    return title_en or kind.title(), detail_en

EN_BULLETS = {
    "新能源行业研究：": ("Green industry research: ", "Researched wind power, geothermal energy, new-energy pipeline networks and SAF-related sectors; summarized policy context, business models and project economics based on project materials."),
    "投资尽调支持：": ("Investment due diligence: ", "Organized corporate, equity, operations, litigation, environmental, social/resettlement and financial due diligence materials for green private-equity projects; identified compliance, operational and exit risks."),
    "投决材料整理：": ("Investment committee materials: ", "Assisted in organizing investment proposals, risk assessment reports, investment committee/general manager/board materials, investment agreements, shareholders' agreements and ancillary agreements; extracted key terms and decision points."),
    "绿色投资项目材料：": ("Green investment materials: ", "Organized investment proposals, risk assessments, internal approval documents and due diligence materials for wind, geothermal and new-energy pipeline projects."),
    "合规与决策支持：": ("Compliance and decision support: ", "Assisted in reviewing project initiation, investment committee, board and internal approval materials; summarized legal, financial, environmental and social due diligence risk points."),
    "绿色基金项目尽调：": ("Green fund due diligence: ", "Organized corporate, transaction, environmental, legal and financial due diligence materials for new-energy investment projects; summarized compliance and transaction risks."),
    "结构性产品文件：": ("Structured products documentation: ", "Supported the Financial Regulation team on structured product documents for foreign banks; assisted in drafting and checking Final Terms and reviewed equity-linked notes, including fixed coupons, autocall mechanics, worst-of basket, observation-date trigger levels, barriers and redemption calculations."),
    "金融监管跟踪：": ("Financial regulatory tracking: ", "Tracked regulatory updates and HKEX announcements for foreign banks; summarized potential implications for structured note issuance, disclosure and compliance arrangements."),
    "条款逻辑核查：": ("Clause logic review: ", "Reviewed highly technical clauses for foreign-bank matters, focusing on non-reliance, no fiduciary, unsecured, no floor, Calculation Agent, selling restrictions and cross-border compliance issues beyond formatting."),
    "合规检索与规则适用：": ("Compliance research and rule application: ", "Researched pharmaceutical regulation, internet-service pre-approval, technology import/export and SEP liability boundaries; distilled approval/filing paths, liability boundaries and practical risk points."),
    "客户咨询 Memo 写作：": ("Client memorandum drafting: ", "Worked directly with a partner on client consultations; independently prepared memorandum drafts organizing regulatory rules into issue lists, applicable rules and risk notes."),
    "交易材料风险核查：": ("Transaction risk review: ", "Reviewed negative information, IP ownership/use restrictions and third-party consent/change-of-control clauses in M&A and investment matters; summarized issues to be confirmed before closing."),
    "证据审阅与整理：": ("Evidence review: ", "Reviewed transaction documents, correspondence and meeting records in a major cross-border investment dispute; identified key evidence on put option exercise, transaction communications and loss issues."),
    "法律检索与文书写作：": ("Legal research and drafting: ", "Worked directly with a partner on two client memoranda; researched pharmaceutical compliance, AI export controls, SEP-related damages and regulatory risk boundaries."),
    "M&A 尽职调查：": ("M&A due diligence: ", "Checked negative information, IP ownership/use restrictions and third-party consent clauses; translated patent materials and summarized risk points."),
    "交易尽调支持：": ("Transaction due diligence: ", "Reviewed negative information, IP ownership, third-party consent and change-of-control clauses in investment-related materials; summarized pre-closing issues and transaction risks."),
    "M&A 文件核查：": ("M&A document review: ", "Assisted in organizing transaction documents, resolutions and closing-stage materials; checked executed versions, exhibit lists and consistency of key clauses."),
    "跨境投融资争议：": ("Cross-border investment dispute: ", "Reviewed transaction communications, meeting records and correspondence; extracted evidence on equity repurchase, investment exit and loss issues."),
    "跨境争议材料分析：": ("Cross-border dispute review: ", "Reviewed transaction communications, meeting records and correspondence; extracted evidence on equity repurchase, investment exit and loss issues."),
    "研究与备忘录：": ("Research and memoranda: ", "Researched pharmaceutical compliance, export controls and SEP liability boundaries; prepared structured memorandum drafts."),
    "合规风险识别：": ("Compliance research: ", "Researched pharmaceutical regulation, technology export controls and internet-service compliance; summarized approval, filing and liability boundaries."),
    "重大项目材料审阅：": ("Project document review: ", "Reviewed cross-border investment dispute materials and organized evidence on transaction communications, exercise notices and loss issues."),
    "商标争议文书：": ("Trademark proceedings: ", "Drafted over 60 submissions for invalidation, opposition and non-use cancellation proceedings, covering prior rights, similarity, bad faith and non-use issues."),
    "近似与恶意分析：": ("Similarity and bad-faith analysis: ", "Compared marks, goods/services subclasses, filing portfolios and prior decisions; organized evidence indexes and procedural materials."),
    "英文客户沟通：": ("English client communication: ", "Prepared over 10 English watch opinions and 40+ client emails summarizing CNIPA decisions, case progress and follow-up strategies."),
    "英文信息整理：": ("English reporting: ", "Prepared English watch opinions and client emails to communicate examination results, case progress and follow-up strategies."),
    "流程沉淀：": ("Workflow handover: ", "Prepared an internship guide covering systems, drafting methods, evidence organization and client reporting workflows."),
    "知识产权风险识别：": ("IP risk analysis: ", "Handled trademark similarity, bad-faith filing, non-use cancellation and well-known mark protection issues; summarized right-stability and evidence risks."),
    "流程规范化：": ("Workflow standardization: ", "Prepared guidance on systems, drafting, evidence filing and handover workflows to improve team efficiency."),
    "民事检察监督：": ("Civil procuratorial supervision: ", "Supported civil litigation, trial-procedure supervision and enforcement supervision matters; mapped procedural paths from first instance to enforcement."),
    "法律文书撰写：": ("Legal document drafting: ", "Assisted in drafting 40+ review reports, litigation-support statements, procuratorial recommendations and protest documents."),
    "事实梳理与文书写作：": ("Fact organization and drafting: ", "Organized civil supervision case materials, procedural timelines and disputed issues; assisted in preparing review reports and legal documents."),
    "公共部门法律实务：": ("Public-sector legal practice: ", "Worked on civil litigation support, trial-procedure supervision and enforcement supervision matters; became familiar with public-sector legal review workflows."),
    "综合文字材料：": ("Formal writing: ", "Assisted in drafting review reports, litigation-support statements, procuratorial recommendations and protest documents, improving fact synthesis and formal writing skills."),
    "商事争议检索：": ("Commercial dispute research: ", "Researched professional lending, early termination, liquidated damages/lost profits, negotiable-instrument disputes and unjust enrichment; prepared research memoranda."),
    "企业尽调支持：": ("Due diligence support: ", "Assisted on a debt restructuring matter by organizing corporate registration, shareholding, operations and litigation information."),
    "债务重组尽调：": ("Debt restructuring due diligence: ", "Organized corporate registration, shareholding, operations and litigation information for a debt restructuring matter, supporting risk assessment."),
    "商事规则研究：": ("Commercial-law research: ", "Researched rules on professional lending, negotiable instruments, unjust enrichment, liquidated damages and lost profits."),
    "企业信息核查：": ("Corporate information review: ", "Organized corporate registration, shareholding, operations and litigation information to support background and risk assessment."),
}

EN_COMPETITIONS = {
    "第23届 Vis East / 第23届贸仲杯 / Moot Shanghai | Oralist": ("Vis (East) Moot / CIETAC Cup / Moot Shanghai | Oralist", "Awards: Second Prize in CIETAC Cup; Team Outstanding Contribution Award; APAC Top 32 in Vis East. Research focus: CISG Article 79 exemption and Article 75 substitute-sale damages."),
    "第十二届史丹森国际环境法模拟法庭竞赛 | Oralist": ("Stetson International Environmental Moot Court Competition | Oralist", "Awards: East Asia Regional Bronze Medalist; advanced to the International Finals. Research focus: CMS treaty exceptions, indigenous subsistence needs, species conservation and sustainable development."),
    "第十六届北外-万慧达知识产权英文模拟法庭竞赛 | Oralist": ("BFSU-Wanhuida IP Moot Court Competition | Oralist", "Awards: National Third Place; Best Respondent Memorial. Research focus: patent non-infringement, prior art defense and absence of damages liability."),
}

EN_RESEARCH = {
    "《替代交易的合理性认定研究》 | 独著": ("Research on the Reasonableness of Substitute Transactions | Sole Author", "Research content: Analyzed the PRC Civil Code standard for substitute-transaction reasonableness by reference to price, timing, quality and judicial practice."),
    "2023全国法科学生模拟立法大赛 | 项目核心成员": ("National Legislative Simulation Competition | Core Member", "AIGC legislation: Researched training-data legality, IP protection, content safety, service-provider obligations and regulatory filing mechanisms; contributed to draft provisions and legislative explanations translating technical risks into workable obligations and liability boundaries."),
}


def english_label(label):
    label = str(label or "").strip()
    if not label:
        return "Work scope: "
    label = label.replace("：", ":")
    if not label.endswith(":"):
        label += ":"
    return label + " "


def normalize_english_resume(candidate, fallback):
    candidate = candidate.get("resume", candidate) if isinstance(candidate, dict) else {}
    if not isinstance(candidate, dict):
        return fallback
    result = {**fallback}
    result["translationMode"] = "model"

    profile = candidate.get("profile") if isinstance(candidate.get("profile"), dict) else {}
    if profile:
        base_profile = fallback.get("profile", {})
        education = []
        for item in profile.get("education", []) if isinstance(profile.get("education"), list) else []:
            if isinstance(item, dict):
                school = str(item.get("school", "")).strip()
                major = str(item.get("major", "")).strip()
                date = str(item.get("date", "")).strip()
                details = item.get("details", [])
            elif isinstance(item, (list, tuple)) and len(item) >= 4:
                school, major, date, details = item[0], item[1], item[2], item[3]
            else:
                continue
            if isinstance(details, str):
                details = [details]
            education.append((str(school).strip(), str(major).strip(), str(date).strip(), [str(d).strip() for d in details if str(d).strip()]))
        result["profile"] = {
            "name": str(profile.get("name") or base_profile.get("name") or "").strip(),
            "contact": str(profile.get("contact") or base_profile.get("contact") or "").strip(),
            "honors": str(profile.get("honors") or base_profile.get("honors") or "").strip(),
            "education": education or base_profile.get("education", []),
        }

    experiences = []
    for exp in candidate.get("experiences", []) if isinstance(candidate.get("experiences"), list) else []:
        if not isinstance(exp, dict):
            continue
        bullets = []
        for item in exp.get("bullets", []) if isinstance(exp.get("bullets"), list) else []:
            if isinstance(item, dict):
                label, text = item.get("label", ""), item.get("text", "")
            elif isinstance(item, (list, tuple)) and len(item) >= 2:
                label, text = item[0], item[1]
            else:
                continue
            if str(text).strip():
                bullets.append((english_label(label), str(text).strip()))
        if bullets:
            experiences.append({
                "org": str(exp.get("org", "")).strip(),
                "role": str(exp.get("role", "")).strip(),
                "date": str(exp.get("date", "")).strip(),
                "bullets": bullets[:3],
            })
    if experiences:
        result["experiences"] = experiences[:5]

    for key in ["competitions", "research"]:
        items = []
        for item in candidate.get(key, []) if isinstance(candidate.get(key), list) else []:
            if isinstance(item, dict):
                items.append((str(item.get("title", "")).strip(), str(item.get("date", "")).strip(), str(item.get("detail", "")).strip()))
            elif isinstance(item, (list, tuple)) and len(item) >= 3:
                items.append((str(item[0]).strip(), str(item[1]).strip(), str(item[2]).strip()))
        if items:
            result[key] = items[:3]
    return result


def translate_resume_local(resume):
    translated = {**resume}
    translated["translationMode"] = "local"
    profile = resume.get("profile", {})
    education = []
    for school, major, date, details in profile.get("education", []):
        school_major = en_school_major(school, major)
        if " | " in school_major:
            school_en, major_en = school_major.split(" | ", 1)
        else:
            school_en, major_en = school_major, ""
        education.append((school_en, major_en, date, [en_detail_text(detail) for detail in details]))
    translated["profile"] = {
        "name": profile.get("enName") or en_name(profile.get("name", "")),
        "contact": profile.get("contact", ""),
        "honors": en_detail_text(profile.get("honors", "")),
        "education": education,
    }

    translated["experiences"] = []
    for exp in resume.get("experiences", []):
        org, role = EN_TRANSLATIONS.get(exp.get("org", ""), (en_org(exp.get("org", "")), en_role(exp.get("role", ""))))
        bullets = []
        for label, text in exp.get("bullets", []):
            en_label, en_text = en_bullet_text(label, text)
            bullets.append((english_label(en_label), en_text))
        translated["experiences"].append({
            "org": org,
            "role": role,
            "date": exp.get("date", ""),
            "bullets": bullets,
        })

    translated["competitions"] = []
    for title, date, detail in resume.get("competitions", []):
        en_title, en_text = en_item(title, detail, "competition")
        translated["competitions"].append((en_title, date, en_text))

    translated["research"] = []
    for title, date, detail in resume.get("research", []):
        en_title, en_text = en_item(title, detail, "research")
        translated["research"].append((en_title, date, en_text))

    translated["opinions"] = list(resume.get("opinions", [])) + [{
        "title": "英文简历生成",
        "body": "英文预览和英文 Word 使用同一套专业英文简历译法；未配置模型 Key 时也可正常生成。",
    }]
    return translated


def translate_resume_with_model(resume):
    fallback = translate_resume_local(resume)
    if not deepseek_is_configured():
        return fallback, "local"
    system = (
        "You are a professional legal resume translator. Polish the provided English resume JSON for a one-page legal/finance resume. "
        "Do not add, invent, exaggerate, delete dates, change schools, change employers, or add clients/project names. "
        "Return JSON only with the same schema: profile, experiences, competitions, research. "
        "Keep dates and numbers exactly. Use concise, polished English resume language."
    )
    payload = {
        "model": get_deepseek_model(),
        "input": [
            {"role": "system", "content": system},
            {"role": "user", "content": json.dumps({"resume": fallback}, ensure_ascii=False)},
        ],
        "temperature": 0.15,
        "response_format": {"type": "json_object"},
    }
    try:
        data = extract_json_object(deepseek_response_text(payload))
        return normalize_english_resume(data, fallback), "model"
    except Exception:
        return fallback, "local"


def build_docx_en(resume):
    doc = Document()
    setup_doc_page(doc, "Times New Roman", "Times New Roman")

    add_name(doc, resume.get("profile", {}).get("enName") or en_name(resume.get("profile", {}).get("name", "")), font="Times New Roman", size=14)
    add_contact(doc, resume.get("profile", {}).get("contact", ""), font="Times New Roman", size=9)

    en_section(doc, "Education & Honors")
    for school, major, date, details in resume.get("profile", {}).get("education", []):
        en_line(doc, en_school_major(school, major), date)
        for detail in details:
            en_detail(doc, en_detail_text(detail))
    honors = resume.get("profile", {}).get("honors", "")
    if honors:
        en_detail(doc, en_detail_text(honors))

    en_section(doc, "Internship Experience")
    for exp in resume.get("experiences", []):
        org, role = EN_TRANSLATIONS.get(exp.get("org", ""), (en_org(exp.get("org", "")), en_role(exp.get("role", ""))))
        en_line(doc, org, exp.get("date", ""))
        en_detail(doc, role, size=9.1)
        for label, text in exp.get("bullets", []):
            en_label, en_text = en_bullet_text(label, text)
            en_bullet(doc, en_label, en_text)

    if resume.get("competitions"):
        en_section(doc, "Competition Experience")
        for title, date, detail in resume.get("competitions", []):
            en_title, en_text = en_item(title, detail, "competition")
            en_line(doc, en_title, date, size=9.2)
            en_bullet(doc, "Awards / Research focus: ", en_text, size=9.0)

    if resume.get("research"):
        en_section(doc, "Research Experience")
        for title, date, detail in resume.get("research", []):
            en_title, en_text = en_item(title, detail, "research")
            en_line(doc, en_title, date, size=9.2)
            en_bullet(doc, "Research content: ", en_text, size=9.0)

    en_section(doc, "Skills & Interests")
    if resume.get("externalParsed"):
        en_detail(doc, en_detail_text(resume.get("profile", {}).get("honors", "")) or "Skills: Legal Research; Microsoft Office; drafting and document review.", size=9.0)
    else:
        en_detail(doc, "Skills & Interests: Please upload your resume or add your skills, languages, certificates and interests in 'Supplementary Materials' to generate personalized content.", size=9.0)

    filename = f"JD定制英文简历_{time.strftime('%Y%m%d_%H%M%S')}.docx"
    output_path = OUTPUTS / filename
    doc.save(output_path)
    return output_path


def extract_docx(payload):
    data = payload.get("data", "")
    if "," in data:
        data = data.split(",", 1)[1]
    raw = base64.b64decode(data)
    tmp = GENERATED / f"upload_{int(time.time() * 1000)}.docx"
    tmp.write_bytes(raw)
    doc = Document(tmp)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append("\t".join(cells))
    return "\n".join(parts)


class Handler(BaseHTTPRequestHandler):
    def _json(self, status, data):
        body = json.dumps(data, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def do_GET(self):
        path = unquote(self.path.split("?", 1)[0])
        if path == "/api/translation-status":
            self._json(200, {"configured": deepseek_is_configured(), "model": "AI", "mode": "model" if deepseek_is_configured() else "local", "provider": "AI"})
            return
        if path == "/":
            path = "/index.html"
        if path.startswith("/download/"):
            file_path = OUTPUTS / Path(path.removeprefix("/download/")).name
            if not file_path.exists():
                self.send_error(404)
                return
            data = file_path.read_bytes()
            self.send_response(200)
            self.send_header("Content-Type", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            self.send_header("Content-Disposition", f"attachment; filename*=UTF-8''{quote(file_path.name)}")
            self.send_header("Content-Length", str(len(data)))
            self.end_headers()
            self.wfile.write(data)
            return
        file_path = STATIC / path.lstrip("/")
        if not file_path.exists() or not file_path.is_file():
            self.send_error(404)
            return
        content_type = "text/html; charset=utf-8"
        if file_path.suffix == ".css":
            content_type = "text/css; charset=utf-8"
        elif file_path.suffix == ".js":
            content_type = "text/javascript; charset=utf-8"
        data = file_path.read_bytes()
        self.send_response(200)
        self.send_header("Content-Type", content_type)
        self.send_header("Content-Length", str(len(data)))
        self.end_headers()
        self.wfile.write(data)

    def do_HEAD(self):
        path = unquote(self.path.split("?", 1)[0])
        if path.startswith("/download/"):
            file_path = OUTPUTS / Path(path.removeprefix("/download/")).name
            if not file_path.exists():
                self.send_error(404)
                return
            self.send_response(200)
            self.send_header("Content-Type", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            self.send_header("Content-Disposition", f"attachment; filename*=UTF-8''{quote(file_path.name)}")
            self.send_header("Content-Length", str(file_path.stat().st_size))
            self.end_headers()
            return
        self.send_response(200)
        self.end_headers()

    def do_POST(self):
        try:
            length = int(self.headers.get("Content-Length", "0"))
            payload = json.loads(self.rfile.read(length).decode("utf-8") or "{}")
            custom = payload.get("custom", "")
            uploads_text = "\n".join(item.get("text", "") for item in payload.get("uploads", []) if isinstance(item, dict))
            if self.path == "/api/translation-key":
                self._json(200, set_runtime_deepseek_config(payload.get("apiKey", ""), payload.get("model", "")))
            elif self.path == "/api/test-connection":
                self._json(200, test_deepseek_connection())
            elif self.path == "/api/translate-en":
                resume = payload.get("resume") if isinstance(payload.get("resume"), dict) else resume_from_payload_or_build(payload, custom, uploads_text)
                en_resume, mode = translate_resume_with_model(resume)
                self._json(200, {"resume": en_resume, "mode": mode})
            elif self.path == "/api/analyze":
                result = build_or_deepseek_rewrite(payload, custom, uploads_text)
                self._json(200, {"resume": result["resume"], "answer": result["answer"], "mode": result["mode"], "safeToExport": result["safeToExport"]})
            elif self.path == "/api/generate":
                resume = resume_from_payload_or_build(payload, custom, uploads_text)
                output = build_docx_template_zh(resume)
                self._json(200, {"file": output.name, "url": f"/download/{output.name}", "resume": resume})
            elif self.path == "/api/generate-en":
                resume = resume_from_payload_or_build(payload, custom, uploads_text)
                en_resume, mode = translate_resume_with_model(resume)
                output = build_docx_template_en(en_resume)
                self._json(200, {"file": output.name, "url": f"/download/{output.name}", "resume": en_resume, "sourceResume": resume, "translationMode": mode})
            elif self.path == "/api/web-research":
                research = run_web_research(payload.get("jd", ""), custom)
                result = build_or_deepseek_rewrite(payload, custom, uploads_text, research)
                self._json(200, {"answer": result["answer"], "resume": result["resume"], "mode": result["mode"], "safeToExport": result["safeToExport"]})
            elif self.path == "/api/extract":
                text = extract_docx(payload)
                self._json(200, {"text": text})
            elif self.path == "/api/upload":
                item = save_uploaded_file(payload)
                self._json(200, item)
            elif self.path == "/api/library":
                library = scan_internship_library(payload.get("jd", ""), custom)
                self._json(200, library)
            else:
                self._json(404, {"error": "not found"})
        except Exception as exc:
            self._json(500, {"error": str(exc)})


def main():
    global RUNTIME_DEEPSEEK_API_KEY, RUNTIME_DEEPSEEK_MODEL
    # 启动时恢复持久化的 DeepSeek key（重启不丢）
    try:
        key_file = LLM_KEY_FILE if LLM_KEY_FILE.exists() else (DEEPSEEK_KEY_FILE if DEEPSEEK_KEY_FILE.exists() else None)
        if key_file:
            saved = key_file.read_text(encoding="utf-8").strip()
            if saved:
                RUNTIME_DEEPSEEK_API_KEY = saved
                if not RUNTIME_DEEPSEEK_MODEL:
                    RUNTIME_DEEPSEEK_MODEL = "deepseek-chat"
    except Exception:
        pass
    port = 8765
    server = ThreadingHTTPServer(("0.0.0.0", port), Handler)
    print(f"监听地址: http://0.0.0.0:{port}（本机访问 http://127.0.0.1:{port}；局域网设备用本机IP访问，如 http://192.168.x.x:{port}）", flush=True)
    server.serve_forever()


if __name__ == "__main__":
    main()
